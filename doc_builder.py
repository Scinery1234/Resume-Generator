import logging
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
from typing import Dict, List
import json
import html

# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# ── Template Definitions ────────────────────────────────────────────────────
#
# layout values:
#   "A" – single-column (Classic, Minimal)
#   "B" – two-column sidebar (Modern, Executive)
#   "C" – full-width header band (Creative)
#
TEMPLATES = {
    "modern": {
        # Layout B – navy sidebar
        "layout":           "B",
        "heading_color":    RGBColor(0x1a, 0x37, 0x5e),
        "rule_color":       RGBColor(0x1a, 0x37, 0x5e),
        "muted_color":      RGBColor(0x44, 0x55, 0x66),
        "body_font":        "Calibri",
        "heading_font":     "Calibri",
        "name_size":        Pt(20),
        "contact_size":     Pt(10),
        "section_size":     Pt(10),
        "body_size":        Pt(10.5),
        "name_align":       WD_ALIGN_PARAGRAPH.LEFT,
        "section_rule":     True,
        # HTML colours
        "html_heading":     "#1a375e",
        "html_rule":        "#1a375e",
        "html_muted":       "#445566",
        "html_name_size":   "20pt",
        "html_body_size":   "10.5pt",
        "html_section_size":"9pt",
        "html_contact_size":"9.5pt",
        "html_accent_bg":   None,
        # Layout B – sidebar palette
        "html_sidebar_bg":          "#1a375e",
        "html_sidebar_text":        "#e8f0f8",
        "html_sidebar_rule":        "#4a7aae",
        "docx_sidebar_bg_hex":      "1A375E",
        "docx_sidebar_text_rgb":    RGBColor(0xe8, 0xf0, 0xf8),
        "docx_sidebar_rule_rgb":    RGBColor(0x4a, 0x7a, 0xae),
    },
    "classic": {
        # Layout A – single column, serif
        "layout":           "A",
        "heading_color":    RGBColor(0x1a, 0x1a, 0x1a),
        "rule_color":       RGBColor(0x33, 0x33, 0x33),
        "muted_color":      RGBColor(0x55, 0x55, 0x55),
        "body_font":        "Georgia",
        "heading_font":     "Georgia",
        "name_size":        Pt(18),
        "contact_size":     Pt(10),
        "section_size":     Pt(10),
        "body_size":        Pt(10.5),
        "name_align":       WD_ALIGN_PARAGRAPH.CENTER,
        "section_rule":     True,
        "html_heading":     "#1a1a1a",
        "html_rule":        "#333333",
        "html_muted":       "#555555",
        "html_name_size":   "18pt",
        "html_body_size":   "10.5pt",
        "html_section_size":"9.5pt",
        "html_contact_size":"10pt",
        "html_accent_bg":   None,
    },
    "creative": {
        # Layout C – full-width purple header band
        "layout":           "C",
        "heading_color":    RGBColor(0x6b, 0x21, 0xa8),
        "rule_color":       RGBColor(0x08, 0x91, 0xb2),
        "muted_color":      RGBColor(0x4b, 0x55, 0x63),
        "body_font":        "Calibri",
        "heading_font":     "Calibri",
        "name_size":        Pt(22),
        "contact_size":     Pt(10),
        "section_size":     Pt(10),
        "body_size":        Pt(10.5),
        "name_align":       WD_ALIGN_PARAGRAPH.LEFT,
        "section_rule":     True,
        "html_heading":     "#6b21a8",
        "html_rule":        "#0891b2",
        "html_muted":       "#4b5563",
        "html_name_size":   "22pt",
        "html_body_size":   "10.5pt",
        "html_section_size":"9pt",
        "html_contact_size":"9.5pt",
        "html_accent_bg":   None,
        # Layout C – header band palette
        "html_header_bg":          "#6b21a8",
        "html_header_text":        "#ffffff",
        "docx_header_bg_hex":      "6B21A8",
        "docx_header_text_rgb":    RGBColor(0xff, 0xff, 0xff),
    },
    "minimal": {
        # Layout A – single column, clean
        "layout":           "A",
        "heading_color":    RGBColor(0x37, 0x41, 0x51),
        "rule_color":       RGBColor(0xd1, 0xd5, 0xdb),
        "muted_color":      RGBColor(0x6b, 0x72, 0x80),
        "body_font":        "Calibri",
        "heading_font":     "Calibri",
        "name_size":        Pt(19),
        "contact_size":     Pt(10),
        "section_size":     Pt(9),
        "body_size":        Pt(10.5),
        "name_align":       WD_ALIGN_PARAGRAPH.LEFT,
        "section_rule":     True,
        "html_heading":     "#374151",
        "html_rule":        "#d1d5db",
        "html_muted":       "#6b7280",
        "html_name_size":   "19pt",
        "html_body_size":   "10.5pt",
        "html_section_size":"8.5pt",
        "html_contact_size":"9.5pt",
        "html_accent_bg":   None,
    },
    "executive": {
        # Layout B – charcoal sidebar with amber-gold rules
        "layout":           "B",
        "heading_color":    RGBColor(0x1c, 0x1c, 0x2e),
        "rule_color":       RGBColor(0xb4, 0x53, 0x09),
        "muted_color":      RGBColor(0x52, 0x52, 0x52),
        "body_font":        "Calibri",
        "heading_font":     "Calibri",
        "name_size":        Pt(22),
        "contact_size":     Pt(10),
        "section_size":     Pt(9.5),
        "body_size":        Pt(10.5),
        "name_align":       WD_ALIGN_PARAGRAPH.LEFT,
        "section_rule":     True,
        "html_heading":     "#1c1c2e",
        "html_rule":        "#b45309",
        "html_muted":       "#525252",
        "html_name_size":   "22pt",
        "html_body_size":   "10.5pt",
        "html_section_size":"9.5pt",
        "html_contact_size":"9.5pt",
        "html_accent_bg":   None,
        # Layout B – sidebar palette
        "html_sidebar_bg":          "#1c1c2e",
        "html_sidebar_text":        "#f0f0f0",
        "html_sidebar_rule":        "#b45309",
        "docx_sidebar_bg_hex":      "1C1C2E",
        "docx_sidebar_text_rgb":    RGBColor(0xf0, 0xf0, 0xf0),
        "docx_sidebar_rule_rgb":    RGBColor(0xb4, 0x53, 0x09),
    },
}

# Available template IDs and metadata (used by the API)
TEMPLATE_LIST = [
    {
        "id":          "modern",
        "name":        "Modern",
        "description": "Navy sidebar with contact & skills panel — polished and professional.",
        "layout":      "B",
    },
    {
        "id":          "classic",
        "name":        "Classic",
        "description": "Traditional serif single-column — timeless and formal.",
        "layout":      "A",
    },
    {
        "id":          "creative",
        "name":        "Creative",
        "description": "Purple header band with teal accents — bold and contemporary.",
        "layout":      "C",
    },
    {
        "id":          "minimal",
        "name":        "Minimal",
        "description": "Light-gray rules, single column — understated and elegant.",
        "layout":      "A",
    },
    {
        "id":          "executive",
        "name":        "Executive",
        "description": "Charcoal sidebar with amber-gold rules — sharp and authoritative.",
        "layout":      "B",
    },
]

# ── Dummy resume for template preview thumbnails ─────────────────────────────
DUMMY_CANDIDATE = {
    "name": "Alex Johnson",
    "contact": {
        "phone":    "0412 345 678",
        "email":    "alex@email.com",
        "location": "Sydney, NSW",
        "linkedin": "linkedin.com/in/alexjohnson",
    },
    "professional_summary": (
        "Results-driven software engineer with 6+ years delivering scalable "
        "web applications. Proven track record leading cross-functional teams "
        "and shipping high-impact features on time and within scope."
    ),
    "key_skills":  ["Python", "TypeScript", "React", "AWS", "Docker", "Agile", "CI/CD"],
    "experience": [
        {
            "title":    "Senior Software Engineer",
            "company":  "Atlassian",
            "location": "Sydney, NSW",
            "dates":    "Jan 2021 – Present",
            "bullets":  [
                "Led architecture of microservices platform serving 2M+ daily active users",
                "Reduced API response time by 40% through caching and query optimisation",
                "Mentored 3 junior engineers, improving team velocity by 25%",
            ],
        },
        {
            "title":    "Software Engineer",
            "company":  "Canva",
            "location": "Sydney, NSW",
            "dates":    "Mar 2018 – Dec 2020",
            "bullets":  [
                "Built real-time collaboration features used by 5M+ users monthly",
                "Implemented automated testing framework, cutting QA time by 35%",
            ],
        },
    ],
    "education": [
        {
            "degree":          "Bachelor of Computer Science",
            "field":           "Software Engineering",
            "institution":     "University of New South Wales",
            "graduation_year": "2018",
        }
    ],
    "certifications":  ["AWS Certified Solutions Architect", "Google Cloud Professional Developer"],
    "awards":          [],
    "technical_skills": ["Python", "JavaScript", "TypeScript", "PostgreSQL", "Redis", "Kubernetes", "Terraform"],
}


# ══════════════════════════════════════════════════════════════════════════════
# Shared helpers
# ══════════════════════════════════════════════════════════════════════════════

def _get_template(template_id: str) -> Dict:
    """Return a template config dict, falling back to 'modern'."""
    return TEMPLATES.get((template_id or "modern").lower(), TEMPLATES["modern"])


def _cm_to_twips(cm: float) -> int:
    """Convert centimetres to Word twips (1 cm ≈ 567.7 twips)."""
    return int(cm * 567.7)


def _set_para_spacing(para, before: int = 0, after: int = 0, line: int = None):
    """Set paragraph spacing in twips (1 pt = 20 twips)."""
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(before * 20))
    spacing.set(qn('w:after'),  str(after  * 20))
    if line is not None:
        spacing.set(qn('w:line'),     str(line))
        spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)


def _set_para_indent(para, left_cm: float = 0, right_cm: float = 0):
    """Set paragraph left/right indentation in cm."""
    pPr = para._p.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    if left_cm:
        ind.set(qn('w:left'),  str(_cm_to_twips(left_cm)))
    if right_cm:
        ind.set(qn('w:right'), str(_cm_to_twips(right_cm)))
    pPr.append(ind)


def _add_rule_to_para(para, color: RGBColor):
    """Add a thin bottom border (horizontal rule) to an existing paragraph."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), f'{color[0]:02X}{color[1]:02X}{color[2]:02X}')
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_horizontal_rule(doc: Document, color: RGBColor = None):
    """Add a standalone thin horizontal-rule paragraph to a Document."""
    if color is None:
        color = RGBColor(0x1a, 0x37, 0x5e)
    para = doc.add_paragraph()
    _set_para_spacing(para, before=0, after=3)
    _add_rule_to_para(para, color)
    return para


def _section_heading(doc: Document, text: str, tmpl: Dict):
    """Add a styled section heading with a rule below (Layout A / C)."""
    para = doc.add_paragraph()
    _set_para_spacing(para, before=10, after=1)
    run = para.add_run(text.upper())
    run.font.name      = tmpl["heading_font"]
    run.font.size      = tmpl["section_size"]
    run.font.bold      = True
    run.font.color.rgb = tmpl["heading_color"]
    if tmpl.get("section_rule", True):
        _add_horizontal_rule(doc, tmpl["rule_color"])


def _set_doc_margins(doc: Document):
    """Set page margins to 2.0 cm top/bottom, 2.2 cm sides (A4)."""
    for section in doc.sections:
        section.page_height   = Cm(29.7)
        section.page_width    = Cm(21.0)
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.2)
        section.right_margin  = Cm(2.2)


def _set_paragraph_shading(para, fill_hex: str):
    """Set a paragraph's background shading colour (fill_hex without '#')."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill_hex.lstrip('#'))
    pPr.append(shd)


def _add_cell_shading(cell, fill_hex: str):
    """Set a table cell's background fill colour (fill_hex without '#')."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill_hex.lstrip('#'))
    tcPr.append(shd)


def _set_cell_width(cell, cm_width: float):
    """Set a table cell's fixed width in cm."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW  = OxmlElement('w:tcW')
    tcW.set(qn('w:w'),    str(_cm_to_twips(cm_width)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)


def _clear_table_borders(table):
    """Remove all visible borders from a table."""
    tbl   = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{name}')
        el.set(qn('w:val'),   'none')
        el.set(qn('w:sz'),    '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _set_table_width(table, cm_width: float):
    """Set a table's total width in cm."""
    tbl   = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'),    str(_cm_to_twips(cm_width)))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)


# ══════════════════════════════════════════════════════════════════════════════
# Layout A – single-column (Classic, Minimal)
# ══════════════════════════════════════════════════════════════════════════════

def _build_word_layout_a(doc: Document, candidate_data: Dict, tmpl: Dict):
    """Build the single-column resume body (Layout A)."""
    _set_doc_margins(doc)

    contact: Dict = candidate_data.get('contact', {})

    # ── Name ────────────────────────────────────────────────────────────────
    name_para = doc.add_paragraph()
    _set_para_spacing(name_para, before=0, after=2)
    name_para.alignment = tmpl["name_align"]
    nr = name_para.add_run(candidate_data.get('name', '').strip().upper())
    nr.font.name      = tmpl["heading_font"]
    nr.font.size      = tmpl["name_size"]
    nr.font.bold      = True
    nr.font.color.rgb = tmpl["heading_color"]

    # ── Contact line ─────────────────────────────────────────────────────────
    contact_parts = []
    if contact.get('phone'):    contact_parts.append(contact['phone'])
    if contact.get('email'):    contact_parts.append(contact['email'])
    if contact.get('location'): contact_parts.append(contact['location'])
    if contact.get('linkedin'): contact_parts.append(contact['linkedin'])

    contact_para = doc.add_paragraph()
    _set_para_spacing(contact_para, before=0, after=4)
    contact_para.alignment = tmpl["name_align"]
    cr = contact_para.add_run('  ·  '.join(contact_parts))
    cr.font.name      = tmpl["body_font"]
    cr.font.size      = tmpl["contact_size"]
    cr.font.color.rgb = tmpl["muted_color"]

    _add_horizontal_rule(doc, tmpl["rule_color"])

    _add_body_sections_a(doc, candidate_data, tmpl)


def _add_body_sections_a(doc: Document, candidate_data: Dict, tmpl: Dict):
    """
    Add the main body sections (Summary → Skills → Experience → Education →
    Certifications → Awards → Tech Skills) to a document or after a header.
    Used by both Layout A and Layout C.
    """
    # ── Professional Summary ─────────────────────────────────────────────────
    summary = candidate_data.get('professional_summary', '').strip()
    if summary:
        _section_heading(doc, 'Professional Summary', tmpl)
        p = doc.add_paragraph()
        _set_para_spacing(p, before=3, after=6, line=276)
        r = p.add_run(summary)
        r.font.name = tmpl["body_font"]
        r.font.size = tmpl["body_size"]

    # ── Key Skills ───────────────────────────────────────────────────────────
    key_skills = candidate_data.get('key_skills', [])
    if key_skills:
        _section_heading(doc, 'Key Skills', tmpl)
        p = doc.add_paragraph()
        _set_para_spacing(p, before=3, after=6)
        r = p.add_run('  ·  '.join(str(s) for s in key_skills))
        r.font.name = tmpl["body_font"]
        r.font.size = tmpl["body_size"]

    # ── Work Experience ──────────────────────────────────────────────────────
    experience = candidate_data.get('experience', [])
    if experience:
        _section_heading(doc, 'Work Experience', tmpl)
        for i, exp in enumerate(experience):
            job_para = doc.add_paragraph()
            _set_para_spacing(job_para, before=4, after=0)
            tr = job_para.add_run(exp.get('title', ''))
            tr.font.name      = tmpl["body_font"]
            tr.font.size      = tmpl["body_size"]
            tr.font.bold      = True
            tr.font.color.rgb = tmpl["heading_color"]

            dates = exp.get('dates', '')
            if dates:
                tab_run = job_para.add_run('\t' + dates)
                tab_run.font.name      = tmpl["body_font"]
                tab_run.font.size      = tmpl["body_size"]
                tab_run.font.bold      = False
                tab_run.font.color.rgb = tmpl["muted_color"]
                pPr  = job_para._p.get_or_add_pPr()
                tabs = OxmlElement('w:tabs')
                tab  = OxmlElement('w:tab')
                tab.set(qn('w:val'), 'right')
                tab.set(qn('w:pos'), '9360')   # ~16.5 cm content width
                tabs.append(tab)
                pPr.append(tabs)

            company  = exp.get('company',  '')
            location = exp.get('location', '')
            sub_parts = [p for p in [company, location] if p]
            if sub_parts:
                sub_para = doc.add_paragraph()
                _set_para_spacing(sub_para, before=0, after=1)
                sr = sub_para.add_run('  |  '.join(sub_parts))
                sr.font.name      = tmpl["body_font"]
                sr.font.size      = tmpl["body_size"]
                sr.font.italic    = True
                sr.font.color.rgb = tmpl["muted_color"]

            bullets     = exp.get('bullets', [])
            description = exp.get('description', '').strip()

            if bullets:
                for bullet in bullets:
                    bp = doc.add_paragraph(style='List Bullet')
                    _set_para_spacing(bp, before=1, after=1)
                    pPr = bp._p.get_or_add_pPr()
                    ind = OxmlElement('w:ind')
                    ind.set(qn('w:left'),    '360')
                    ind.set(qn('w:hanging'), '180')
                    pPr.append(ind)
                    br = bp.add_run(str(bullet))
                    br.font.name = tmpl["body_font"]
                    br.font.size = tmpl["body_size"]
            elif description:
                dp = doc.add_paragraph()
                _set_para_spacing(dp, before=1, after=1, line=276)
                dr = dp.add_run(description)
                dr.font.name = tmpl["body_font"]
                dr.font.size = tmpl["body_size"]

            if i < len(experience) - 1:
                sp = doc.add_paragraph()
                _set_para_spacing(sp, before=0, after=3)

    # ── Education ────────────────────────────────────────────────────────────
    education = candidate_data.get('education', [])
    if education:
        _section_heading(doc, 'Education', tmpl)
        for edu in education:
            degree   = edu.get('degree', '')
            field    = edu.get('field',  '')
            grad     = edu.get('graduation_year', '')
            deg_text = f"{degree}{' — ' + field if field else ''}"

            edu_para = doc.add_paragraph()
            _set_para_spacing(edu_para, before=4, after=0)
            dr = edu_para.add_run(deg_text)
            dr.font.name      = tmpl["body_font"]
            dr.font.size      = tmpl["body_size"]
            dr.font.bold      = True
            dr.font.color.rgb = tmpl["heading_color"]
            if grad:
                tab_run = edu_para.add_run('\t' + grad)
                tab_run.font.name      = tmpl["body_font"]
                tab_run.font.size      = tmpl["body_size"]
                tab_run.font.bold      = False
                tab_run.font.color.rgb = tmpl["muted_color"]
                pPr    = edu_para._p.get_or_add_pPr()
                tabs   = OxmlElement('w:tabs')
                tab_el = OxmlElement('w:tab')
                tab_el.set(qn('w:val'), 'right')
                tab_el.set(qn('w:pos'), '9360')
                tabs.append(tab_el)
                pPr.append(tabs)

            inst = edu.get('institution', '')
            if inst:
                ip = doc.add_paragraph()
                _set_para_spacing(ip, before=0, after=4)
                ir = ip.add_run(inst)
                ir.font.name      = tmpl["body_font"]
                ir.font.size      = tmpl["body_size"]
                ir.font.italic    = True
                ir.font.color.rgb = tmpl["muted_color"]

    # ── Certifications ───────────────────────────────────────────────────────
    certs = candidate_data.get('certifications', [])
    if certs:
        _section_heading(doc, 'Certifications', tmpl)
        for cert in certs:
            cp = doc.add_paragraph(style='List Bullet')
            _set_para_spacing(cp, before=1, after=1)
            cr = cp.add_run(str(cert))
            cr.font.name = tmpl["body_font"]
            cr.font.size = tmpl["body_size"]

    # ── Awards ───────────────────────────────────────────────────────────────
    awards = candidate_data.get('awards', [])
    if awards:
        _section_heading(doc, 'Awards & Recognition', tmpl)
        for award in awards:
            ap = doc.add_paragraph(style='List Bullet')
            _set_para_spacing(ap, before=1, after=1)
            ar = ap.add_run(str(award))
            ar.font.name = tmpl["body_font"]
            ar.font.size = tmpl["body_size"]

    # ── Technical Skills ─────────────────────────────────────────────────────
    tech_skills = candidate_data.get('technical_skills', [])
    if tech_skills:
        _section_heading(doc, 'Technical Skills', tmpl)
        p = doc.add_paragraph()
        _set_para_spacing(p, before=3, after=6)
        r = p.add_run('  ·  '.join(str(s) for s in tech_skills))
        r.font.name = tmpl["body_font"]
        r.font.size = tmpl["body_size"]


# ══════════════════════════════════════════════════════════════════════════════
# Layout B – two-column sidebar (Modern, Executive)
# ══════════════════════════════════════════════════════════════════════════════

_SIDEBAR_CM  = 6.5    # sidebar column width
_MAIN_CM     = 14.5   # main column width  (total = 21 cm)
_SB_PAD_L    = 1.5    # sidebar left padding (cm)
_SB_PAD_R    = 0.8    # sidebar right padding (cm)
_MN_PAD_L    = 1.8    # main left padding (cm)
_MN_PAD_R    = 1.5    # main right padding (cm)
# Right-aligned tab stop inside the main cell, measured from cell left edge:
#   cell_width - right_pad = 14.5 - 1.5 = 13.0 cm
_MN_TAB_POS  = _cm_to_twips(_MAIN_CM - _MN_PAD_R)


def _sb_heading(cell, text: str, tmpl: Dict):
    """Add a sidebar section heading + coloured rule to a table cell."""
    hp = cell.add_paragraph()
    _set_para_spacing(hp, before=14, after=1)
    _set_para_indent(hp, _SB_PAD_L, _SB_PAD_R)
    hr = hp.add_run(text.upper())
    hr.font.name      = tmpl["heading_font"]
    hr.font.size      = Pt(8.5)
    hr.font.bold      = True
    hr.font.color.rgb = tmpl["docx_sidebar_text_rgb"]

    rp = cell.add_paragraph()
    _set_para_spacing(rp, before=0, after=5)
    _set_para_indent(rp, _SB_PAD_L, _SB_PAD_R)
    _add_rule_to_para(rp, tmpl["docx_sidebar_rule_rgb"])


def _sb_text(cell, text: str, tmpl: Dict, italic: bool = False,
             font_size: Pt = None, before: int = 2, after: int = 2):
    """Add a plain text paragraph to the sidebar cell."""
    p = cell.add_paragraph()
    _set_para_spacing(p, before=before, after=after)
    _set_para_indent(p, _SB_PAD_L, _SB_PAD_R)
    r = p.add_run(text)
    r.font.name      = tmpl["body_font"]
    r.font.size      = font_size or Pt(9.5)
    r.font.italic    = italic
    r.font.color.rgb = tmpl["docx_sidebar_text_rgb"]


def _mn_heading(cell, text: str, tmpl: Dict):
    """Add a main-column section heading + rule to a table cell."""
    hp = cell.add_paragraph()
    _set_para_spacing(hp, before=12, after=1)
    _set_para_indent(hp, _MN_PAD_L, _MN_PAD_R)
    hr = hp.add_run(text.upper())
    hr.font.name      = tmpl["heading_font"]
    hr.font.size      = tmpl["section_size"]
    hr.font.bold      = True
    hr.font.color.rgb = tmpl["heading_color"]

    rp = cell.add_paragraph()
    _set_para_spacing(rp, before=0, after=4)
    _set_para_indent(rp, _MN_PAD_L, _MN_PAD_R)
    _add_rule_to_para(rp, tmpl["rule_color"])


def _build_word_layout_b(doc: Document, candidate_data: Dict, tmpl: Dict):
    """
    Build the two-column sidebar layout (Layout B).
    Left column: dark background with contact / key-skills / education / certs.
    Right column: white background with summary / experience / awards / tech-skills.
    """
    # Zero outer margins — the table fills the full page width
    for sec in doc.sections:
        sec.page_height   = Cm(29.7)
        sec.page_width    = Cm(21.0)
        sec.top_margin    = Cm(0)
        sec.bottom_margin = Cm(0.5)
        sec.left_margin   = Cm(0)
        sec.right_margin  = Cm(0)

    table = doc.add_table(rows=1, cols=2)
    _clear_table_borders(table)
    _set_table_width(table, 21.0)

    sb = table.cell(0, 0)   # sidebar
    mn = table.cell(0, 1)   # main

    _set_cell_width(sb, _SIDEBAR_CM)
    _set_cell_width(mn, _MAIN_CM)
    _add_cell_shading(sb, tmpl["docx_sidebar_bg_hex"])

    # ── Sidebar ───────────────────────────────────────────────────────────────
    contact: Dict = candidate_data.get('contact', {})

    # Top spacer (reuse the default empty paragraph)
    sp0 = sb.paragraphs[0]
    _set_para_spacing(sp0, before=30, after=0)

    # Name
    np_ = sb.add_paragraph()
    _set_para_spacing(np_, before=0, after=10)
    _set_para_indent(np_, _SB_PAD_L, _SB_PAD_R)
    nr = np_.add_run(candidate_data.get('name', '').strip().upper())
    nr.font.name      = tmpl["heading_font"]
    nr.font.size      = Pt(16)
    nr.font.bold      = True
    nr.font.color.rgb = tmpl["docx_sidebar_text_rgb"]

    # CONTACT
    _sb_heading(sb, 'Contact', tmpl)
    for item in [contact.get('phone'), contact.get('email'),
                 contact.get('location'), contact.get('linkedin')]:
        if item:
            _sb_text(sb, item, tmpl)

    # KEY SKILLS
    key_skills = candidate_data.get('key_skills', [])
    if key_skills:
        _sb_heading(sb, 'Key Skills', tmpl)
        for skill in key_skills:
            _sb_text(sb, f'  {skill}', tmpl)

    # EDUCATION
    education = candidate_data.get('education', [])
    if education:
        _sb_heading(sb, 'Education', tmpl)
        for edu in education:
            degree   = edu.get('degree', '')
            field    = edu.get('field',  '')
            deg_text = f"{degree}{' — ' + field if field else ''}"
            _sb_text(sb, deg_text, tmpl, font_size=Pt(9))
            if edu.get('institution'):
                _sb_text(sb, edu['institution'], tmpl, italic=True, font_size=Pt(8.5))
            if edu.get('graduation_year'):
                _sb_text(sb, edu['graduation_year'], tmpl, font_size=Pt(8.5))

    # CERTIFICATIONS
    certs = candidate_data.get('certifications', [])
    if certs:
        _sb_heading(sb, 'Certifications', tmpl)
        for cert in certs:
            _sb_text(sb, f'  {cert}', tmpl, font_size=Pt(9))

    # Bottom spacer
    _sb_text(sb, '', tmpl, before=0, after=20)

    # ── Main column ───────────────────────────────────────────────────────────
    # Top spacer (reuse the default empty paragraph)
    sp1 = mn.paragraphs[0]
    _set_para_spacing(sp1, before=30, after=0)

    # PROFESSIONAL SUMMARY
    summary = candidate_data.get('professional_summary', '').strip()
    if summary:
        _mn_heading(mn, 'Professional Summary', tmpl)
        p = mn.add_paragraph()
        _set_para_spacing(p, before=3, after=6, line=276)
        _set_para_indent(p, _MN_PAD_L, _MN_PAD_R)
        r = p.add_run(summary)
        r.font.name = tmpl["body_font"]
        r.font.size = tmpl["body_size"]

    # WORK EXPERIENCE
    experience = candidate_data.get('experience', [])
    if experience:
        _mn_heading(mn, 'Work Experience', tmpl)
        for i, exp in enumerate(experience):
            # Job title + right-aligned dates
            jp = mn.add_paragraph()
            _set_para_spacing(jp, before=4, after=0)
            _set_para_indent(jp, _MN_PAD_L, _MN_PAD_R)
            tr = jp.add_run(exp.get('title', ''))
            tr.font.name      = tmpl["body_font"]
            tr.font.size      = tmpl["body_size"]
            tr.font.bold      = True
            tr.font.color.rgb = tmpl["heading_color"]

            dates = exp.get('dates', '')
            if dates:
                tab_run = jp.add_run('\t' + dates)
                tab_run.font.name      = tmpl["body_font"]
                tab_run.font.size      = tmpl["body_size"]
                tab_run.font.bold      = False
                tab_run.font.color.rgb = tmpl["muted_color"]
                pPr  = jp._p.get_or_add_pPr()
                tabs = OxmlElement('w:tabs')
                tab  = OxmlElement('w:tab')
                tab.set(qn('w:val'), 'right')
                tab.set(qn('w:pos'), str(_MN_TAB_POS))
                tabs.append(tab)
                pPr.append(tabs)

            company  = exp.get('company',  '')
            location = exp.get('location', '')
            sub_parts = [x for x in [company, location] if x]
            if sub_parts:
                sub = mn.add_paragraph()
                _set_para_spacing(sub, before=0, after=1)
                _set_para_indent(sub, _MN_PAD_L, _MN_PAD_R)
                sr = sub.add_run('  |  '.join(sub_parts))
                sr.font.name      = tmpl["body_font"]
                sr.font.size      = tmpl["body_size"]
                sr.font.italic    = True
                sr.font.color.rgb = tmpl["muted_color"]

            bullets     = exp.get('bullets', [])
            description = exp.get('description', '').strip()

            if bullets:
                for bullet in bullets:
                    bp = mn.add_paragraph()
                    _set_para_spacing(bp, before=1, after=1)
                    _set_para_indent(bp, _MN_PAD_L + 0.35, _MN_PAD_R)
                    pPr = bp._p.get_or_add_pPr()
                    ind = OxmlElement('w:ind')
                    ind.set(qn('w:hanging'), '200')
                    pPr.append(ind)
                    br = bp.add_run('•  ' + str(bullet))
                    br.font.name = tmpl["body_font"]
                    br.font.size = tmpl["body_size"]
            elif description:
                dp = mn.add_paragraph()
                _set_para_spacing(dp, before=1, after=1, line=276)
                _set_para_indent(dp, _MN_PAD_L, _MN_PAD_R)
                dr = dp.add_run(description)
                dr.font.name = tmpl["body_font"]
                dr.font.size = tmpl["body_size"]

            if i < len(experience) - 1:
                gap = mn.add_paragraph()
                _set_para_spacing(gap, before=0, after=4)

    # AWARDS
    awards = candidate_data.get('awards', [])
    if awards:
        _mn_heading(mn, 'Awards & Recognition', tmpl)
        for award in awards:
            ap = mn.add_paragraph()
            _set_para_spacing(ap, before=1, after=1)
            _set_para_indent(ap, _MN_PAD_L + 0.35, _MN_PAD_R)
            ar = ap.add_run('•  ' + str(award))
            ar.font.name = tmpl["body_font"]
            ar.font.size = tmpl["body_size"]

    # TECHNICAL SKILLS
    tech_skills = candidate_data.get('technical_skills', [])
    if tech_skills:
        _mn_heading(mn, 'Technical Skills', tmpl)
        p = mn.add_paragraph()
        _set_para_spacing(p, before=3, after=6)
        _set_para_indent(p, _MN_PAD_L, _MN_PAD_R)
        r = p.add_run('  ·  '.join(str(s) for s in tech_skills))
        r.font.name = tmpl["body_font"]
        r.font.size = tmpl["body_size"]

    # Bottom spacer
    gap = mn.add_paragraph()
    _set_para_spacing(gap, before=0, after=20)


# ══════════════════════════════════════════════════════════════════════════════
# Layout C – full-width header band (Creative)
# ══════════════════════════════════════════════════════════════════════════════

def _build_word_layout_c(doc: Document, candidate_data: Dict, tmpl: Dict):
    """
    Build the header-band layout (Layout C).
    A shaded header at the top (name + contact on coloured background) is
    followed by the standard single-column body sections.
    """
    _set_doc_margins(doc)

    contact: Dict = candidate_data.get('contact', {})
    contact_parts = []
    if contact.get('phone'):    contact_parts.append(contact['phone'])
    if contact.get('email'):    contact_parts.append(contact['email'])
    if contact.get('location'): contact_parts.append(contact['location'])
    if contact.get('linkedin'): contact_parts.append(contact['linkedin'])

    header_hex  = tmpl["docx_header_bg_hex"]
    header_rgb  = tmpl["docx_header_text_rgb"]

    # ── Name (shaded) ────────────────────────────────────────────────────────
    name_para = doc.add_paragraph()
    _set_para_spacing(name_para, before=14, after=2)
    _set_paragraph_shading(name_para, header_hex)
    nr = name_para.add_run(candidate_data.get('name', '').strip().upper())
    nr.font.name      = tmpl["heading_font"]
    nr.font.size      = tmpl["name_size"]
    nr.font.bold      = True
    nr.font.color.rgb = header_rgb

    # ── Contact (shaded) ─────────────────────────────────────────────────────
    contact_para = doc.add_paragraph()
    _set_para_spacing(contact_para, before=0, after=14)
    _set_paragraph_shading(contact_para, header_hex)
    cr = contact_para.add_run('  ·  '.join(contact_parts))
    cr.font.name      = tmpl["body_font"]
    cr.font.size      = tmpl["contact_size"]
    cr.font.color.rgb = header_rgb

    # Rule separating header from body
    _add_horizontal_rule(doc, tmpl["rule_color"])

    # ── Body (same as Layout A) ───────────────────────────────────────────────
    _add_body_sections_a(doc, candidate_data, tmpl)


# ══════════════════════════════════════════════════════════════════════════════
# HTML helpers (shared between layouts)
# ══════════════════════════════════════════════════════════════════════════════

def _html_body_sections(candidate_data: Dict, tmpl: Dict) -> List[str]:
    """
    Build the list of HTML section strings that are common to Layout A and C.
    Returns a list of '<div class="section">…</div>' strings.
    """
    e = html.escape
    parts: List[str] = []

    def section(title: str, body: str) -> str:
        return (
            f'<div class="section">'
            f'  <div class="section-heading">{e(title).upper()}</div>'
            f'  <div class="section-rule"></div>'
            f'  {body}'
            f'</div>'
        )

    # Professional Summary
    summary = candidate_data.get('professional_summary', '').strip()
    if summary:
        parts.append(section('Professional Summary',
            f'<p class="body-text">{e(summary)}</p>'))

    # Key Skills
    key_skills = candidate_data.get('key_skills', [])
    if key_skills:
        parts.append(section('Key Skills',
            f'<p class="body-text">{"  ·  ".join(e(str(s)) for s in key_skills)}</p>'))

    # Work Experience
    experience = candidate_data.get('experience', [])
    if experience:
        exp_html = ''
        for exp in experience:
            title    = e(exp.get('title',    ''))
            company  = e(exp.get('company',  ''))
            location = e(exp.get('location', ''))
            dates    = e(exp.get('dates',    ''))
            desc     = e(exp.get('description', '').strip())
            bullets  = exp.get('bullets', [])
            sub      = '  |  '.join(x for x in [company, location] if x)
            exp_html += (
                f'<div class="exp-entry">'
                f'  <div class="exp-header">'
                f'    <span class="exp-title">{title}</span>'
                f'    <span class="exp-dates">{dates}</span>'
                f'  </div>'
                f'  <div class="exp-sub">{sub}</div>'
            )
            if bullets:
                exp_html += '<ul class="exp-bullets">'
                for b in bullets:
                    exp_html += f'<li>{e(str(b))}</li>'
                exp_html += '</ul>'
            elif desc:
                exp_html += f'<p class="body-text">{desc}</p>'
            exp_html += '</div>'
        parts.append(section('Work Experience', exp_html))

    # Education
    education = candidate_data.get('education', [])
    if education:
        edu_html = ''
        for edu in education:
            degree   = e(edu.get('degree', ''))
            field    = e(edu.get('field',  ''))
            inst     = e(edu.get('institution', ''))
            grad     = e(edu.get('graduation_year', ''))
            deg_text = f"{degree}{' — ' + field if field else ''}"
            edu_html += (
                f'<div class="edu-entry">'
                f'  <div class="edu-header">'
                f'    <span class="edu-degree">{deg_text}</span>'
                f'    <span class="edu-year">{grad}</span>'
                f'  </div>'
                f'  <div class="edu-inst">{inst}</div>'
                f'</div>'
            )
        parts.append(section('Education', edu_html))

    # Certifications
    certs = candidate_data.get('certifications', [])
    if certs:
        c_html = '<ul class="exp-bullets">' + ''.join(f'<li>{e(str(c))}</li>' for c in certs) + '</ul>'
        parts.append(section('Certifications', c_html))

    # Awards
    awards = candidate_data.get('awards', [])
    if awards:
        a_html = '<ul class="exp-bullets">' + ''.join(f'<li>{e(str(a))}</li>' for a in awards) + '</ul>'
        parts.append(section('Awards & Recognition', a_html))

    # Technical Skills
    tech_skills = candidate_data.get('technical_skills', [])
    if tech_skills:
        parts.append(section('Technical Skills',
            f'<p class="body-text">{"  ·  ".join(e(str(s)) for s in tech_skills)}</p>'))

    return parts


def _base_css(tmpl: Dict, font_family: str) -> str:
    """Return CSS rules common to section headings, experience, and education."""
    return f'''
  .section {{ margin-bottom: 16px; }}
  .section-heading {{
    font-size: {tmpl["html_section_size"]};
    font-weight: 700;
    color: {tmpl["html_heading"]};
    letter-spacing: 0.08em;
    margin-bottom: 2px;
    margin-top: 4px;
  }}
  .section-rule {{
    border-top: 1px solid {tmpl["html_rule"]};
    margin-bottom: 8px;
  }}
  .body-text {{
    font-size: {tmpl["html_body_size"]};
    line-height: 1.4;
    color: #222;
  }}
  .exp-entry {{ margin-bottom: 10px; }}
  .exp-header {{
    display: flex;
    justify-content: space-between;
    align-items: baseline;
  }}
  .exp-title  {{ font-weight: 700; font-size: {tmpl["html_body_size"]}; color: {tmpl["html_heading"]}; }}
  .exp-dates  {{ font-size: 9.5pt; color: {tmpl["html_muted"]}; white-space: nowrap; }}
  .exp-sub    {{ font-style: italic; color: {tmpl["html_muted"]}; font-size: 9.5pt; margin-bottom: 4px; }}
  .exp-bullets {{
    margin-left: 1.1em;
    padding-left: 0.5em;
    font-size: {tmpl["html_body_size"]};
    line-height: 1.5;
  }}
  .exp-bullets li {{ margin-bottom: 2px; }}
  .edu-entry {{ margin-bottom: 8px; }}
  .edu-header {{
    display: flex;
    justify-content: space-between;
    align-items: baseline;
  }}
  .edu-degree {{ font-weight: 700; font-size: {tmpl["html_body_size"]}; color: {tmpl["html_heading"]}; }}
  .edu-year   {{ font-size: 9.5pt; color: {tmpl["html_muted"]}; }}
  .edu-inst   {{ font-style: italic; color: {tmpl["html_muted"]}; font-size: 9.5pt; }}
'''


# ══════════════════════════════════════════════════════════════════════════════
# HTML Layout A – single-column
# ══════════════════════════════════════════════════════════════════════════════

def _build_html_layout_a(candidate_data: Dict, tmpl: Dict) -> str:
    """Return self-contained HTML for the single-column layout (A)."""
    e = html.escape
    contact: Dict = candidate_data.get('contact', {})

    contact_parts = []
    if contact.get('phone'):    contact_parts.append(e(contact['phone']))
    if contact.get('email'):
        em = e(contact['email'])
        contact_parts.append(f'<a href="mailto:{em}">{em}</a>')
    if contact.get('location'): contact_parts.append(e(contact['location']))
    if contact.get('linkedin'):
        li = e(contact['linkedin'])
        contact_parts.append(f'<a href="https://{li}" target="_blank">{li}</a>')

    font_family = (
        "Georgia, 'Times New Roman', serif"
        if tmpl["body_font"] == "Georgia"
        else "Calibri, 'Segoe UI', Arial, sans-serif"
    )
    name_align = "left" if tmpl["name_align"] == WD_ALIGN_PARAGRAPH.LEFT else "center"

    body_sections = ''.join(_html_body_sections(candidate_data, tmpl))
    name_html    = e(candidate_data.get('name', '').strip().upper())
    contact_html = '  ·  '.join(contact_parts)

    return f'''<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8"/>
<meta name="viewport" content="width=device-width, initial-scale=1.0"/>
<title>Resume Preview</title>
<style>
  *, *::before, *::after {{ box-sizing: border-box; margin: 0; padding: 0; }}
  body {{
    font-family: {font_family};
    font-size: {tmpl["html_body_size"]};
    color: #222;
    background: #fff;
  }}
  .page {{
    width: 21cm;
    min-height: 29.7cm;
    padding: 2cm 2.2cm;
    margin: 0 auto;
    background: #fff;
  }}
  .resume-name {{
    font-size: {tmpl["html_name_size"]};
    font-weight: 700;
    color: {tmpl["html_heading"]};
    text-align: {name_align};
    letter-spacing: 0.04em;
    margin-bottom: 4px;
  }}
  .resume-contact {{
    font-size: {tmpl["html_contact_size"]};
    color: {tmpl["html_muted"]};
    text-align: {name_align};
    margin-bottom: 10px;
  }}
  .resume-contact a {{ color: {tmpl["html_muted"]}; text-decoration: none; }}
  .header-rule {{
    border: none;
    border-top: 1.5px solid {tmpl["html_rule"]};
    margin-bottom: 14px;
  }}
  {_base_css(tmpl, font_family)}
</style>
</head>
<body>
<div class="page">
  <div class="resume-name">{name_html}</div>
  <div class="resume-contact">{contact_html}</div>
  <hr class="header-rule"/>
  {body_sections}
</div>
</body>
</html>'''


# ══════════════════════════════════════════════════════════════════════════════
# HTML Layout B – two-column sidebar
# ══════════════════════════════════════════════════════════════════════════════

def _build_html_layout_b(candidate_data: Dict, tmpl: Dict) -> str:
    """Return self-contained HTML for the two-column sidebar layout (B)."""
    e = html.escape
    contact: Dict = candidate_data.get('contact', {})

    sb_bg   = tmpl["html_sidebar_bg"]
    sb_text = tmpl["html_sidebar_text"]
    sb_rule = tmpl["html_sidebar_rule"]

    font_family = "Calibri, 'Segoe UI', Arial, sans-serif"

    # ── Build sidebar content ─────────────────────────────────────────────────
    def sb_section(label: str, items_html: str) -> str:
        return (
            f'<div class="sb-section">'
            f'  <div class="sb-heading">{e(label).upper()}</div>'
            f'  <div class="sb-rule"></div>'
            f'  {items_html}'
            f'</div>'
        )

    sb_parts = []

    # Contact
    contact_items = ''
    for val in [contact.get('phone'), contact.get('email'),
                contact.get('location'), contact.get('linkedin')]:
        if val:
            contact_items += f'<div class="sb-item">{e(val)}</div>'
    if contact_items:
        sb_parts.append(sb_section('Contact', contact_items))

    # Key Skills
    key_skills = candidate_data.get('key_skills', [])
    if key_skills:
        skills_html = ''.join(f'<div class="sb-item">  {e(str(s))}</div>' for s in key_skills)
        sb_parts.append(sb_section('Key Skills', skills_html))

    # Education
    education = candidate_data.get('education', [])
    if education:
        edu_html = ''
        for edu in education:
            degree   = e(edu.get('degree', ''))
            field    = e(edu.get('field',  ''))
            inst     = e(edu.get('institution', ''))
            grad     = e(edu.get('graduation_year', ''))
            deg_text = f"{degree}{' — ' + field if field else ''}"
            edu_html += f'<div class="sb-item">{deg_text}</div>'
            if inst:
                edu_html += f'<div class="sb-item-muted">{inst}</div>'
            if grad:
                edu_html += f'<div class="sb-item" style="font-size:8.5pt;">{grad}</div>'
        sb_parts.append(sb_section('Education', edu_html))

    # Certifications
    certs = candidate_data.get('certifications', [])
    if certs:
        c_html = ''.join(f'<div class="sb-item">  {e(str(c))}</div>' for c in certs)
        sb_parts.append(sb_section('Certifications', c_html))

    sidebar_html = f'<div class="sidebar-name">{e(candidate_data.get("name","").strip().upper())}</div>' + ''.join(sb_parts)

    # ── Build main-column content ─────────────────────────────────────────────
    main_parts: List[str] = []
    section_css_classes = ''  # reuse _html_body_sections but only for main

    def main_section(title: str, body: str) -> str:
        return (
            f'<div class="section">'
            f'  <div class="section-heading">{e(title).upper()}</div>'
            f'  <div class="section-rule"></div>'
            f'  {body}'
            f'</div>'
        )

    # Role title block — pushes Professional Summary below the sidebar name
    experience = candidate_data.get('experience', [])
    current_role = ''
    if experience:
        first_exp = experience[0]
        role_parts = [first_exp.get('title', '')]
        if first_exp.get('company'):
            role_parts.append(first_exp['company'])
        current_role = '  |  '.join(r for r in role_parts if r)
    if current_role:
        main_parts.append(f'<div class="main-role">{e(current_role)}</div>')

    summary = candidate_data.get('professional_summary', '').strip()
    if summary:
        main_parts.append(main_section('Professional Summary',
            f'<p class="body-text">{e(summary)}</p>'))

    if experience:
        exp_html = ''
        for exp in experience:
            title    = e(exp.get('title',    ''))
            company  = e(exp.get('company',  ''))
            location = e(exp.get('location', ''))
            dates    = e(exp.get('dates',    ''))
            desc     = e(exp.get('description', '').strip())
            bullets  = exp.get('bullets', [])
            sub      = '  |  '.join(x for x in [company, location] if x)
            exp_html += (
                f'<div class="exp-entry">'
                f'  <div class="exp-header">'
                f'    <span class="exp-title">{title}</span>'
                f'    <span class="exp-dates">{dates}</span>'
                f'  </div>'
                f'  <div class="exp-sub">{sub}</div>'
            )
            if bullets:
                exp_html += '<ul class="exp-bullets">'
                for b in bullets:
                    exp_html += f'<li>{e(str(b))}</li>'
                exp_html += '</ul>'
            elif desc:
                exp_html += f'<p class="body-text">{desc}</p>'
            exp_html += '</div>'
        main_parts.append(main_section('Work Experience', exp_html))

    awards = candidate_data.get('awards', [])
    if awards:
        a_html = '<ul class="exp-bullets">' + ''.join(f'<li>{e(str(a))}</li>' for a in awards) + '</ul>'
        main_parts.append(main_section('Awards & Recognition', a_html))

    tech_skills = candidate_data.get('technical_skills', [])
    if tech_skills:
        main_parts.append(main_section('Technical Skills',
            f'<p class="body-text">{"  ·  ".join(e(str(s)) for s in tech_skills)}</p>'))

    main_html = ''.join(main_parts)

    return f'''<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8"/>
<meta name="viewport" content="width=device-width, initial-scale=1.0"/>
<title>Resume Preview</title>
<style>
  *, *::before, *::after {{ box-sizing: border-box; margin: 0; padding: 0; }}
  body {{
    font-family: {font_family};
    font-size: {tmpl["html_body_size"]};
    color: #222;
    background: #fff;
  }}
  .page {{
    width: 21cm;
    min-height: 29.7cm;
    display: flex;
    background: #fff;
  }}
  /* ── Sidebar ── */
  .sidebar {{
    width: 31%;
    min-height: 29.7cm;
    background: {sb_bg};
    padding: 1.8cm 1.4cm 2cm 1.5cm;
    flex-shrink: 0;
  }}
  .sidebar-name {{
    font-size: 15pt;
    font-weight: 700;
    color: {sb_text};
    letter-spacing: 0.04em;
    margin-bottom: 14px;
    line-height: 1.2;
    word-break: break-word;
  }}
  .sb-section {{ margin-bottom: 14px; }}
  .sb-heading {{
    font-size: 8pt;
    font-weight: 700;
    color: {sb_text};
    letter-spacing: 0.1em;
    margin-bottom: 3px;
  }}
  .sb-rule {{
    border-top: 1px solid {sb_rule};
    margin-bottom: 6px;
  }}
  .sb-item {{
    font-size: 9pt;
    color: {sb_text};
    margin-bottom: 3px;
    line-height: 1.4;
    word-break: break-all;
  }}
  .sb-item-muted {{
    font-size: 8.5pt;
    color: {sb_text};
    opacity: 0.8;
    font-style: italic;
    margin-bottom: 2px;
  }}
  /* ── Main column ── */
  .main {{
    flex: 1;
    padding: 1.8cm 1.8cm 2cm 1.8cm;
    min-width: 0;
  }}
  .main-role {{
    font-size: 11pt;
    font-weight: 600;
    color: {tmpl["html_heading"]};
    letter-spacing: 0.02em;
    margin-bottom: 18px;
    padding-bottom: 10px;
    border-bottom: 1px solid {tmpl["html_rule"]};
  }}
  {_base_css(tmpl, font_family)}
</style>
</head>
<body>
<div class="page">
  <div class="sidebar">
    {sidebar_html}
  </div>
  <div class="main">
    {main_html}
  </div>
</div>
</body>
</html>'''


# ══════════════════════════════════════════════════════════════════════════════
# HTML Layout C – full-width header band
# ══════════════════════════════════════════════════════════════════════════════

def _build_html_layout_c(candidate_data: Dict, tmpl: Dict) -> str:
    """Return self-contained HTML for the header-band layout (C)."""
    e = html.escape
    contact: Dict = candidate_data.get('contact', {})

    hdr_bg   = tmpl["html_header_bg"]
    hdr_text = tmpl["html_header_text"]

    contact_parts = []
    if contact.get('phone'):    contact_parts.append(e(contact['phone']))
    if contact.get('email'):
        em = e(contact['email'])
        contact_parts.append(f'<a href="mailto:{em}" style="color:{hdr_text};text-decoration:none;">{em}</a>')
    if contact.get('location'): contact_parts.append(e(contact['location']))
    if contact.get('linkedin'):
        li = e(contact['linkedin'])
        contact_parts.append(f'<a href="https://{li}" target="_blank" style="color:{hdr_text};text-decoration:none;">{li}</a>')

    font_family = "Calibri, 'Segoe UI', Arial, sans-serif"

    body_sections = ''.join(_html_body_sections(candidate_data, tmpl))
    name_html    = e(candidate_data.get('name', '').strip().upper())
    contact_html = '  ·  '.join(contact_parts)

    return f'''<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8"/>
<meta name="viewport" content="width=device-width, initial-scale=1.0"/>
<title>Resume Preview</title>
<style>
  *, *::before, *::after {{ box-sizing: border-box; margin: 0; padding: 0; }}
  body {{
    font-family: {font_family};
    font-size: {tmpl["html_body_size"]};
    color: #222;
    background: #fff;
  }}
  .page {{
    width: 21cm;
    min-height: 29.7cm;
    background: #fff;
  }}
  /* ── Header band ── */
  .header-band {{
    background: {hdr_bg};
    padding: 1.5cm 2.2cm 1.2cm;
    text-align: left;
  }}
  .hb-name {{
    font-size: {tmpl["html_name_size"]};
    font-weight: 700;
    color: {hdr_text};
    letter-spacing: 0.04em;
    margin-bottom: 5px;
    line-height: 1.2;
  }}
  .hb-contact {{
    font-size: {tmpl["html_contact_size"]};
    color: rgba(255,255,255,0.85);
  }}
  .hb-contact a {{ color: rgba(255,255,255,0.85); text-decoration: none; }}
  /* ── Body ── */
  .page-body {{
    padding: 1.2cm 2.2cm 2cm;
  }}
  {_base_css(tmpl, font_family)}
</style>
</head>
<body>
<div class="page">
  <div class="header-band">
    <div class="hb-name">{name_html}</div>
    <div class="hb-contact">{contact_html}</div>
  </div>
  <div class="page-body">
    {body_sections}
  </div>
</div>
</body>
</html>'''


# ══════════════════════════════════════════════════════════════════════════════
# ResumeBuilder
# ══════════════════════════════════════════════════════════════════════════════

class ResumeBuilder:
    def __init__(self):
        self.resume_data: Dict = {}

    # ── Legacy collect methods (kept for compatibility) ──────────────────────
    def collect_personal_info(self, name, contact_info):
        self.resume_data['name'] = name
        self.resume_data['contact_info'] = contact_info

    def collect_experience(self, experience):
        self.resume_data['experience'] = experience

    def collect_education(self, education):
        self.resume_data['education'] = education

    def build_resume_text(self) -> str:
        """Plain-text fallback."""
        resume = f"{self.resume_data.get('name', '')}\n"
        contact = self.resume_data.get('contact_info', {})
        if isinstance(contact, dict):
            resume += f"{contact.get('email', '')} | {contact.get('phone', '')} | {contact.get('location', '')}\n\n"
        resume += 'WORK EXPERIENCE\n'
        for exp in self.resume_data.get('experience', []):
            if isinstance(exp, dict):
                resume += f"  {exp.get('title', '')} at {exp.get('company', '')} ({exp.get('dates', '')})\n"
        resume += '\nEDUCATION\n'
        for edu in self.resume_data.get('education', []):
            if isinstance(edu, dict):
                resume += f"  {edu.get('degree', '')} from {edu.get('institution', '')}\n"
        return resume

    # ── Main Word document builder ───────────────────────────────────────────
    def build_word_document(self, output_path: str, candidate_data: Dict,
                            template_id: str = "modern") -> str:
        """
        Build a professionally formatted resume as a .docx file.
        Dispatches to the correct layout builder based on template.

        Args:
            output_path:    Where to save the .docx file.
            candidate_data: Resume data dictionary.
            template_id:    One of "modern", "classic", "creative", "minimal", "executive".
        """
        tmpl   = _get_template(template_id)
        layout = tmpl.get("layout", "A")

        doc = Document()
        normal = doc.styles['Normal']
        normal.font.name = tmpl["body_font"]
        normal.font.size = tmpl["body_size"]

        if layout == "B":
            _build_word_layout_b(doc, candidate_data, tmpl)
        elif layout == "C":
            _build_word_layout_c(doc, candidate_data, tmpl)
        else:
            _build_word_layout_a(doc, candidate_data, tmpl)

        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        logger.info(f'Resume saved → {output_path}  (template: {template_id}, layout: {layout})')
        return output_path

    # ── HTML preview builder ─────────────────────────────────────────────────
    def build_html_preview(self, candidate_data: Dict,
                           template_id: str = "modern") -> str:
        """
        Return a self-contained HTML string that visually matches the .docx.
        Dispatches to the correct layout renderer based on template.

        Args:
            candidate_data: Resume data dictionary.
            template_id:    One of "modern", "classic", "creative", "minimal", "executive".
        """
        tmpl   = _get_template(template_id)
        layout = tmpl.get("layout", "A")

        if layout == "B":
            return _build_html_layout_b(candidate_data, tmpl)
        elif layout == "C":
            return _build_html_layout_c(candidate_data, tmpl)
        else:
            return _build_html_layout_a(candidate_data, tmpl)


# ══════════════════════════════════════════════════════════════════════════════
# Smoke test
# ══════════════════════════════════════════════════════════════════════════════

if __name__ == '__main__':
    sample = {
        "name": "Jane Smith",
        "contact": {
            "phone":    "0412 345 678",
            "email":    "jane.smith@email.com",
            "location": "Melbourne, VIC",
            "linkedin": "linkedin.com/in/janesmith",
        },
        "professional_summary": (
            "Experienced software engineer with 7+ years delivering scalable "
            "web applications for the Australian financial services sector."
        ),
        "key_skills": ["Python", "React", "AWS", "Agile", "CI/CD"],
        "experience": [
            {
                "title":    "Senior Software Engineer",
                "company":  "ANZ Bank",
                "location": "Melbourne, VIC",
                "dates":    "Jan 2020 – Present",
                "description": "",
                "bullets": [
                    "Led a team of 5 engineers delivering a $2M digital banking platform",
                    "Reduced deployment time by 60% through CI/CD pipeline optimisation",
                ],
            }
        ],
        "education": [
            {
                "degree":          "Bachelor of Computer Science",
                "field":           "Software Engineering",
                "institution":     "University of Melbourne",
                "graduation_year": "2016",
            }
        ],
        "certifications": ["AWS Certified Solutions Architect"],
        "awards":          [],
        "technical_skills": ["Python", "JavaScript", "PostgreSQL", "Docker", "Kubernetes"],
    }
    builder = ResumeBuilder()
    for tid in ["modern", "classic", "creative", "minimal", "executive"]:
        builder.build_word_document(f'/tmp/test_resume_{tid}.docx', sample, tid)
        html_out = builder.build_html_preview(sample, tid)
        with open(f'/tmp/test_resume_{tid}.html', 'w') as f:
            f.write(html_out)
        print(f"Smoke test complete: {tid}")
