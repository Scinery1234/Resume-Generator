"""
Unit tests for doc_builder.py – ResumeBuilder class.
Tests cover: Word document generation, HTML preview generation, all 3 layouts,
edge cases, and template configuration.
"""
import os
import tempfile
import pytest
from pathlib import Path
from docx import Document

from doc_builder import ResumeBuilder, TEMPLATES, TEMPLATE_LIST, _get_template


# ── Fixtures ────────────────────────────────────────────────────────────────

FULL_CANDIDATE = {
    "name": "Jane Smith",
    "contact": {
        "phone":    "0412 345 678",
        "email":    "jane.smith@email.com",
        "location": "Melbourne, VIC",
        "linkedin": "linkedin.com/in/janesmith",
    },
    "professional_summary": (
        "Experienced software engineer with 7+ years delivering scalable "
        "web applications for the Australian financial services sector."
    ),
    "key_skills": ["Python", "React", "AWS"],
    "experience": [
        {
            "title":    "Senior Software Engineer",
            "company":  "ANZ Bank",
            "location": "Melbourne, VIC",
            "dates":    "Jan 2020 – Present",
            "description": "",
            "bullets": [
                "Led a team of 5 engineers",
                "Reduced deployment time by 60%",
            ],
        },
        {
            "title":    "Software Developer",
            "company":  "Startup Co",
            "location": "Sydney, NSW",
            "dates":    "Mar 2017 – Dec 2019",
            "description": "Developed REST APIs and front-end features.",
            "bullets": [],
        },
    ],
    "education": [
        {
            "degree":          "Bachelor of Computer Science",
            "field":           "Software Engineering",
            "institution":     "University of Melbourne",
            "graduation_year": "2016",
        }
    ],
    "certifications":  ["AWS Certified Solutions Architect"],
    "awards":          ["Employee of the Year 2022"],
    "technical_skills": ["Python", "JavaScript", "PostgreSQL", "Docker"],
}

MINIMAL_CANDIDATE = {
    "name": "Bob Jones",
    "contact": {
        "email":    "bob@example.com",
        "phone":    "0400000000",
        "location": "Brisbane, QLD",
    },
    "professional_summary": "Minimal candidate for testing.",
    "key_skills":       [],
    "experience":       [],
    "education":        [],
    "certifications":   [],
    "awards":           [],
    "technical_skills": [],
}


@pytest.fixture
def builder():
    return ResumeBuilder()


@pytest.fixture
def tmp_docx(tmp_path):
    return str(tmp_path / "test_resume.docx")


# ── Helper ───────────────────────────────────────────────────────────────────

def _get_all_text(doc: Document) -> str:
    """
    Extract all text from a Document, including paragraphs inside table cells.
    Layout B stores content in table cells which doc.paragraphs does not expose.
    """
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(p.text for p in cell.paragraphs)
    return " ".join(parts)


# ══════════════════════════════════════════════════════════════════════════════
# Word document tests  (default template = "modern", which is Layout B)
# We use classic for the Layout A-specific structural tests.
# ══════════════════════════════════════════════════════════════════════════════

class TestBuildWordDocument:
    def test_creates_file(self, builder, tmp_docx):
        builder.build_word_document(tmp_docx, FULL_CANDIDATE)
        assert Path(tmp_docx).exists()
        assert Path(tmp_docx).stat().st_size > 0

    def test_minimal_candidate_creates_file(self, builder, tmp_docx):
        builder.build_word_document(tmp_docx, MINIMAL_CANDIDATE)
        assert Path(tmp_docx).exists()

    def test_returns_output_path(self, builder, tmp_docx):
        result = builder.build_word_document(tmp_docx, FULL_CANDIDATE)
        assert result == tmp_docx

    def test_docx_contains_candidate_name(self, builder, tmp_docx):
        builder.build_word_document(tmp_docx, FULL_CANDIDATE)
        doc = Document(tmp_docx)
        assert "JANE SMITH" in _get_all_text(doc)

    def test_docx_contains_contact_info(self, builder, tmp_docx):
        builder.build_word_document(tmp_docx, FULL_CANDIDATE)
        doc = Document(tmp_docx)
        full_text = _get_all_text(doc)
        assert "jane.smith@email.com" in full_text
        assert "0412 345 678" in full_text
        assert "Melbourne, VIC" in full_text

    def test_docx_contains_professional_summary(self, builder, tmp_docx):
        builder.build_word_document(tmp_docx, FULL_CANDIDATE)
        doc = Document(tmp_docx)
        assert "financial services sector" in _get_all_text(doc)

    def test_docx_contains_key_skills(self, builder, tmp_docx):
        builder.build_word_document(tmp_docx, FULL_CANDIDATE)
        doc = Document(tmp_docx)
        full_text = _get_all_text(doc)
        assert "Python" in full_text
        assert "React" in full_text

    def test_docx_contains_experience_title(self, builder, tmp_docx):
        builder.build_word_document(tmp_docx, FULL_CANDIDATE)
        doc = Document(tmp_docx)
        full_text = _get_all_text(doc)
        assert "Senior Software Engineer" in full_text
        assert "ANZ Bank" in full_text

    def test_docx_contains_experience_bullets(self, builder, tmp_docx):
        builder.build_word_document(tmp_docx, FULL_CANDIDATE)
        doc = Document(tmp_docx)
        assert "Led a team of 5 engineers" in _get_all_text(doc)

    def test_docx_contains_experience_description(self, builder, tmp_docx):
        """Second experience entry uses description (no bullets)."""
        builder.build_word_document(tmp_docx, FULL_CANDIDATE)
        doc = Document(tmp_docx)
        assert "Developed REST APIs" in _get_all_text(doc)

    def test_docx_contains_education(self, builder, tmp_docx):
        builder.build_word_document(tmp_docx, FULL_CANDIDATE)
        doc = Document(tmp_docx)
        full_text = _get_all_text(doc)
        assert "Bachelor of Computer Science" in full_text
        assert "University of Melbourne" in full_text

    def test_docx_contains_certifications(self, builder, tmp_docx):
        builder.build_word_document(tmp_docx, FULL_CANDIDATE)
        doc = Document(tmp_docx)
        assert "AWS Certified Solutions Architect" in _get_all_text(doc)

    def test_docx_contains_awards(self, builder, tmp_docx):
        builder.build_word_document(tmp_docx, FULL_CANDIDATE)
        doc = Document(tmp_docx)
        assert "Employee of the Year 2022" in _get_all_text(doc)

    def test_docx_contains_technical_skills(self, builder, tmp_docx):
        builder.build_word_document(tmp_docx, FULL_CANDIDATE)
        doc = Document(tmp_docx)
        assert "PostgreSQL" in _get_all_text(doc)

    def test_missing_optional_sections_skipped(self, builder, tmp_docx):
        """When no certifications/awards, those headings should not appear."""
        builder.build_word_document(tmp_docx, MINIMAL_CANDIDATE)
        doc = Document(tmp_docx)
        full_text = _get_all_text(doc).upper()
        assert "CERTIFICATIONS" not in full_text
        assert "AWARDS" not in full_text

    def test_creates_parent_directory(self, builder, tmp_path):
        nested_path = str(tmp_path / "nested" / "dir" / "resume.docx")
        builder.build_word_document(nested_path, MINIMAL_CANDIDATE)
        assert Path(nested_path).exists()

    def test_linkedin_included_in_contact(self, builder, tmp_docx):
        builder.build_word_document(tmp_docx, FULL_CANDIDATE)
        doc = Document(tmp_docx)
        assert "linkedin.com/in/janesmith" in _get_all_text(doc)

    def test_no_linkedin_still_works(self, builder, tmp_docx):
        builder.build_word_document(tmp_docx, MINIMAL_CANDIDATE)
        assert Path(tmp_docx).exists()


# ── HTML preview tests ───────────────────────────────────────────────────────

class TestBuildHtmlPreview:
    def test_returns_string(self, builder):
        result = builder.build_html_preview(FULL_CANDIDATE)
        assert isinstance(result, str)

    def test_contains_html_structure(self, builder):
        result = builder.build_html_preview(FULL_CANDIDATE)
        assert "<!DOCTYPE html>" in result
        assert "<body>" in result
        assert "</html>" in result

    def test_contains_candidate_name_uppercased(self, builder):
        result = builder.build_html_preview(FULL_CANDIDATE)
        assert "JANE SMITH" in result

    def test_contains_contact_email(self, builder):
        result = builder.build_html_preview(FULL_CANDIDATE)
        assert "jane.smith@email.com" in result

    def test_contains_phone_and_location(self, builder):
        result = builder.build_html_preview(FULL_CANDIDATE)
        assert "0412 345 678" in result
        assert "Melbourne, VIC" in result

    def test_contains_linkedin_link(self, builder):
        result = builder.build_html_preview(FULL_CANDIDATE)
        assert "linkedin.com/in/janesmith" in result

    def test_contains_professional_summary(self, builder):
        result = builder.build_html_preview(FULL_CANDIDATE)
        assert "financial services sector" in result

    def test_contains_key_skills(self, builder):
        result = builder.build_html_preview(FULL_CANDIDATE)
        assert "Python" in result
        assert "React" in result

    def test_contains_experience_entries(self, builder):
        result = builder.build_html_preview(FULL_CANDIDATE)
        assert "Senior Software Engineer" in result
        assert "ANZ Bank" in result

    def test_contains_bullets(self, builder):
        result = builder.build_html_preview(FULL_CANDIDATE)
        assert "Led a team of 5 engineers" in result

    def test_contains_description_fallback(self, builder):
        result = builder.build_html_preview(FULL_CANDIDATE)
        assert "Developed REST APIs" in result

    def test_contains_education(self, builder):
        result = builder.build_html_preview(FULL_CANDIDATE)
        assert "Bachelor of Computer Science" in result
        assert "University of Melbourne" in result

    def test_contains_certifications(self, builder):
        result = builder.build_html_preview(FULL_CANDIDATE)
        assert "AWS Certified Solutions Architect" in result

    def test_contains_awards(self, builder):
        result = builder.build_html_preview(FULL_CANDIDATE)
        assert "Employee of the Year 2022" in result

    def test_contains_technical_skills(self, builder):
        result = builder.build_html_preview(FULL_CANDIDATE)
        assert "PostgreSQL" in result

    def test_html_escapes_special_chars(self, builder):
        candidate = dict(MINIMAL_CANDIDATE)
        candidate["name"] = '<Script>alert("xss")</Script>'
        result = builder.build_html_preview(candidate)
        assert "<Script>" not in result
        assert "<SCRIPT>" not in result
        assert "&lt;SCRIPT&gt;" in result

    def test_minimal_candidate_produces_valid_html(self, builder):
        result = builder.build_html_preview(MINIMAL_CANDIDATE)
        assert "BOB JONES" in result

    def test_optional_sections_absent_when_empty(self, builder):
        result = builder.build_html_preview(MINIMAL_CANDIDATE).upper()
        assert "CERTIFICATIONS" not in result
        assert "AWARDS" not in result
        assert "KEY SKILLS" not in result


# ── Legacy methods ──────────────────────────────────────────────────────────

class TestLegacyMethods:
    def test_collect_personal_info(self, builder):
        builder.collect_personal_info("Alice", {"email": "a@b.com"})
        assert builder.resume_data["name"] == "Alice"

    def test_collect_experience(self, builder):
        exp = [{"title": "Dev", "company": "Acme"}]
        builder.collect_experience(exp)
        assert builder.resume_data["experience"] == exp

    def test_build_resume_text(self, builder):
        builder.collect_personal_info("Alice", {"email": "a@b.com", "phone": "0400", "location": "Sydney"})
        builder.collect_experience([{"title": "Dev", "company": "Acme", "dates": "2020"}])
        builder.collect_education([{"degree": "B.Sc.", "institution": "UNSW"}])
        text = builder.build_resume_text()
        assert "Alice" in text
        assert "Dev" in text


# ── Template configuration tests ─────────────────────────────────────────────

TEMPLATE_IDS = ["modern", "classic", "creative", "minimal", "executive"]


class TestTemplateConfig:
    def test_all_five_templates_defined(self):
        for tid in TEMPLATE_IDS:
            assert tid in TEMPLATES

    def test_each_template_has_required_keys(self):
        required = {
            "heading_color", "rule_color", "muted_color",
            "body_font", "heading_font",
            "name_size", "contact_size", "section_size", "body_size",
            "name_align",
            "html_heading", "html_rule", "html_muted",
            "html_name_size", "html_body_size",
            "html_section_size", "html_contact_size",
            "layout",
        }
        for tid in TEMPLATE_IDS:
            missing = required - TEMPLATES[tid].keys()
            assert not missing, f"Template '{tid}' missing keys: {missing}"

    def test_layout_a_templates(self):
        assert TEMPLATES["classic"]["layout"] == "A"
        assert TEMPLATES["minimal"]["layout"] == "A"

    def test_layout_b_templates(self):
        assert TEMPLATES["modern"]["layout"] == "B"
        assert TEMPLATES["executive"]["layout"] == "B"

    def test_layout_c_templates(self):
        assert TEMPLATES["creative"]["layout"] == "C"

    def test_layout_b_templates_have_sidebar_keys(self):
        sidebar_keys = {
            "html_sidebar_bg", "html_sidebar_text", "html_sidebar_rule",
            "docx_sidebar_bg_hex", "docx_sidebar_text_rgb", "docx_sidebar_rule_rgb",
        }
        for tid in ["modern", "executive"]:
            missing = sidebar_keys - TEMPLATES[tid].keys()
            assert not missing, f"Layout B template '{tid}' missing sidebar keys: {missing}"

    def test_layout_c_template_has_header_keys(self):
        header_keys = {
            "html_header_bg", "html_header_text",
            "docx_header_bg_hex", "docx_header_text_rgb",
        }
        missing = header_keys - TEMPLATES["creative"].keys()
        assert not missing, f"Layout C template 'creative' missing header keys: {missing}"

    def test_get_template_returns_modern_by_default(self):
        assert _get_template(None) is TEMPLATES["modern"]
        assert _get_template("") is TEMPLATES["modern"]

    def test_get_template_unknown_falls_back_to_modern(self):
        assert _get_template("unknown-template") is TEMPLATES["modern"]

    def test_get_template_case_insensitive(self):
        assert _get_template("Modern") is TEMPLATES["modern"]
        assert _get_template("CLASSIC") is TEMPLATES["classic"]

    def test_template_list_has_five_entries(self):
        assert len(TEMPLATE_LIST) == 5

    def test_template_list_ids_match_templates(self):
        list_ids = {t["id"] for t in TEMPLATE_LIST}
        assert list_ids == set(TEMPLATE_IDS)

    def test_template_list_entries_have_required_fields(self):
        for entry in TEMPLATE_LIST:
            assert "id" in entry
            assert "name" in entry
            assert "description" in entry
            assert "layout" in entry

    def test_templates_have_distinct_heading_colors(self):
        colors = [TEMPLATES[tid]["html_heading"] for tid in TEMPLATE_IDS]
        assert len(set(colors)) == len(TEMPLATE_IDS), \
            "All templates should have distinct heading colours"


# ── Per-template Word document tests ─────────────────────────────────────────

class TestBuildWordDocumentWithTemplates:
    @pytest.mark.parametrize("template_id", TEMPLATE_IDS)
    def test_all_templates_create_file(self, builder, tmp_path, template_id):
        path = str(tmp_path / f"resume_{template_id}.docx")
        builder.build_word_document(path, FULL_CANDIDATE, template_id=template_id)
        assert Path(path).exists()
        assert Path(path).stat().st_size > 0

    @pytest.mark.parametrize("template_id", TEMPLATE_IDS)
    def test_all_templates_contain_name(self, builder, tmp_path, template_id):
        path = str(tmp_path / f"resume_{template_id}.docx")
        builder.build_word_document(path, FULL_CANDIDATE, template_id=template_id)
        doc = Document(path)
        assert "JANE SMITH" in _get_all_text(doc)

    @pytest.mark.parametrize("template_id", TEMPLATE_IDS)
    def test_all_templates_contain_experience(self, builder, tmp_path, template_id):
        path = str(tmp_path / f"resume_{template_id}.docx")
        builder.build_word_document(path, FULL_CANDIDATE, template_id=template_id)
        doc = Document(path)
        assert "Senior Software Engineer" in _get_all_text(doc)

    @pytest.mark.parametrize("template_id", TEMPLATE_IDS)
    def test_all_templates_contain_contact(self, builder, tmp_path, template_id):
        path = str(tmp_path / f"resume_{template_id}.docx")
        builder.build_word_document(path, FULL_CANDIDATE, template_id=template_id)
        doc = Document(path)
        assert "jane.smith@email.com" in _get_all_text(doc)

    def test_unknown_template_falls_back_gracefully(self, builder, tmp_docx):
        """Unknown template IDs should fall back to 'modern' without raising."""
        builder.build_word_document(tmp_docx, MINIMAL_CANDIDATE, template_id="nonexistent")
        assert Path(tmp_docx).exists()

    def test_classic_uses_georgia_font(self, builder, tmp_path):
        path = str(tmp_path / "classic.docx")
        builder.build_word_document(path, FULL_CANDIDATE, template_id="classic")
        doc = Document(path)
        assert doc.styles["Normal"].font.name == "Georgia"

    def test_layout_a_name_paragraph_alignment(self, builder, tmp_path):
        """Classic (Layout A) centres the name; minimal (Layout A) left-aligns it."""
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        path_c = str(tmp_path / "classic.docx")
        builder.build_word_document(path_c, FULL_CANDIDATE, template_id="classic")
        doc_c = Document(path_c)
        assert doc_c.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER

        path_m = str(tmp_path / "minimal.docx")
        builder.build_word_document(path_m, FULL_CANDIDATE, template_id="minimal")
        doc_m = Document(path_m)
        assert doc_m.paragraphs[0].alignment in (WD_ALIGN_PARAGRAPH.LEFT, None)

    def test_layout_b_uses_table(self, builder, tmp_path):
        """Modern and Executive (Layout B) render their content inside a table."""
        for tid in ["modern", "executive"]:
            path = str(tmp_path / f"{tid}.docx")
            builder.build_word_document(path, FULL_CANDIDATE, template_id=tid)
            doc = Document(path)
            assert len(doc.tables) >= 1, f"{tid} should have at least one table"

    def test_layout_b_sidebar_contains_contact(self, builder, tmp_path):
        """Sidebar cell (column 0) should contain contact information."""
        path = str(tmp_path / "modern.docx")
        builder.build_word_document(path, FULL_CANDIDATE, template_id="modern")
        doc = Document(path)
        sidebar_text = " ".join(p.text for p in doc.tables[0].cell(0, 0).paragraphs)
        assert "jane.smith@email.com" in sidebar_text

    def test_layout_b_main_contains_experience(self, builder, tmp_path):
        """Main cell (column 1) should contain work experience."""
        path = str(tmp_path / "modern.docx")
        builder.build_word_document(path, FULL_CANDIDATE, template_id="modern")
        doc = Document(path)
        main_text = " ".join(p.text for p in doc.tables[0].cell(0, 1).paragraphs)
        assert "Senior Software Engineer" in main_text

    def test_layout_c_name_is_in_first_paragraph(self, builder, tmp_path):
        """Creative (Layout C) renders the name in the first paragraph (shaded header)."""
        path = str(tmp_path / "creative.docx")
        builder.build_word_document(path, FULL_CANDIDATE, template_id="creative")
        doc = Document(path)
        # First paragraph holds the shaded name
        assert "JANE SMITH" in doc.paragraphs[0].text


# ── Per-template HTML tests ───────────────────────────────────────────────────

class TestBuildHtmlPreviewWithTemplates:
    @pytest.mark.parametrize("template_id", TEMPLATE_IDS)
    def test_all_templates_produce_valid_html(self, builder, template_id):
        result = builder.build_html_preview(FULL_CANDIDATE, template_id=template_id)
        assert "<!DOCTYPE html>" in result
        assert "JANE SMITH" in result

    @pytest.mark.parametrize("template_id", TEMPLATE_IDS)
    def test_all_templates_contain_heading_color(self, builder, template_id):
        result = builder.build_html_preview(FULL_CANDIDATE, template_id=template_id)
        expected_color = TEMPLATES[template_id]["html_heading"]
        assert expected_color in result

    def test_modern_html_uses_sans_font(self, builder):
        result = builder.build_html_preview(FULL_CANDIDATE, template_id="modern")
        assert "Calibri" in result

    def test_classic_html_uses_serif_font(self, builder):
        result = builder.build_html_preview(FULL_CANDIDATE, template_id="classic")
        assert "Georgia" in result

    def test_creative_html_has_purple_heading(self, builder):
        result = builder.build_html_preview(FULL_CANDIDATE, template_id="creative")
        assert "#6b21a8" in result

    def test_minimal_html_has_light_rule(self, builder):
        result = builder.build_html_preview(FULL_CANDIDATE, template_id="minimal")
        assert "#d1d5db" in result

    def test_unknown_template_falls_back_gracefully(self, builder):
        result = builder.build_html_preview(MINIMAL_CANDIDATE, template_id="unknown")
        assert "<!DOCTYPE html>" in result

    # ── Layout A structural tests ───────────────────────────────────────────
    def test_classic_html_name_center_aligned(self, builder):
        result = builder.build_html_preview(FULL_CANDIDATE, template_id="classic")
        assert "text-align: center" in result

    def test_minimal_html_name_left_aligned(self, builder):
        result = builder.build_html_preview(FULL_CANDIDATE, template_id="minimal")
        assert "text-align: left" in result

    # ── Layout B structural tests ───────────────────────────────────────────
    def test_layout_b_html_has_sidebar_div(self, builder):
        for tid in ["modern", "executive"]:
            result = builder.build_html_preview(FULL_CANDIDATE, template_id=tid)
            assert 'class="sidebar"' in result, f"{tid} should have sidebar div"

    def test_layout_b_html_has_main_div(self, builder):
        for tid in ["modern", "executive"]:
            result = builder.build_html_preview(FULL_CANDIDATE, template_id=tid)
            assert 'class="main"' in result, f"{tid} should have main div"

    def test_layout_b_html_sidebar_contains_contact(self, builder):
        result = builder.build_html_preview(FULL_CANDIDATE, template_id="modern")
        sidebar_start = result.find('class="sidebar"')
        main_start    = result.find('class="main"')
        sidebar_block = result[sidebar_start:main_start]
        assert "jane.smith@email.com" in sidebar_block

    def test_layout_b_html_main_contains_experience(self, builder):
        result = builder.build_html_preview(FULL_CANDIDATE, template_id="executive")
        main_start = result.find('class="main"')
        main_block  = result[main_start:]
        assert "Senior Software Engineer" in main_block

    def test_modern_html_has_navy_sidebar(self, builder):
        result = builder.build_html_preview(FULL_CANDIDATE, template_id="modern")
        assert "#1a375e" in result   # sidebar background

    def test_executive_html_has_amber_rule(self, builder):
        result = builder.build_html_preview(FULL_CANDIDATE, template_id="executive")
        assert "#b45309" in result   # amber-gold sidebar rule

    # ── Layout C structural tests ───────────────────────────────────────────
    def test_layout_c_html_has_header_band(self, builder):
        result = builder.build_html_preview(FULL_CANDIDATE, template_id="creative")
        assert 'class="header-band"' in result

    def test_layout_c_html_has_page_body(self, builder):
        result = builder.build_html_preview(FULL_CANDIDATE, template_id="creative")
        assert 'class="page-body"' in result

    def test_layout_c_html_name_in_header_band(self, builder):
        result = builder.build_html_preview(FULL_CANDIDATE, template_id="creative")
        band_start = result.find('class="header-band"')
        body_start = result.find('class="page-body"')
        band_block = result[band_start:body_start]
        assert "JANE SMITH" in band_block

    def test_creative_html_has_purple_header_bg(self, builder):
        result = builder.build_html_preview(FULL_CANDIDATE, template_id="creative")
        assert "#6b21a8" in result   # header band background

    def test_layout_c_html_has_teal_rule(self, builder):
        result = builder.build_html_preview(FULL_CANDIDATE, template_id="creative")
        assert "#0891b2" in result   # section rule colour
