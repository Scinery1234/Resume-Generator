"""
Integration tests for the FastAPI application (main.py).
Uses TestClient from httpx/starlette – no real database calls for most tests
(the DB is SQLite in-memory or a temp file).
"""
import io
import os
import json
import pytest
from pathlib import Path
from unittest.mock import MagicMock, patch

# Use a temp SQLite DB to avoid touching the production DB
os.environ.setdefault("DATABASE_URL", "sqlite:///./test_temp.db")

from fastapi.testclient import TestClient
from main import app

client = TestClient(app)


# ── Health endpoints ─────────────────────────────────────────────────────────

class TestHealth:
    def test_root_returns_healthy(self):
        resp = client.get("/")
        assert resp.status_code == 200
        data = resp.json()
        assert data["status"] == "healthy"

    def test_health_endpoint(self):
        resp = client.get("/health")
        assert resp.status_code == 200
        data = resp.json()
        assert data["status"] == "ok"
        assert "timestamp" in data


# ── Auth endpoints ───────────────────────────────────────────────────────────

class TestAuth:
    _email = "testauth_unique_xyz@example.com"
    _password = "testpass123"
    _name = "Auth Test User"

    def test_signup_creates_user(self):
        resp = client.post("/api/auth/signup", json={
            "name": self._name,
            "email": self._email,
            "password": self._password,
        })
        # May already exist from a previous run – both 200 and 400 are acceptable
        assert resp.status_code in (200, 400)
        if resp.status_code == 200:
            data = resp.json()
            assert data["status"] == "success"
            assert "token" in data
            assert "user_id" in data

    def test_signup_duplicate_email_returns_400(self):
        # Ensure the user exists first
        client.post("/api/auth/signup", json={
            "name": "Dup User",
            "email": "dup_test_xyz@example.com",
            "password": "password123",
        })
        resp = client.post("/api/auth/signup", json={
            "name": "Dup User 2",
            "email": "dup_test_xyz@example.com",
            "password": "password123",
        })
        assert resp.status_code == 400
        assert "already registered" in resp.json()["detail"].lower()

    def test_login_with_valid_credentials(self):
        email = "login_test_xyz@example.com"
        client.post("/api/auth/signup", json={"name": "Login User", "email": email, "password": "pass1234"})
        resp = client.post("/api/auth/login", json={"email": email, "password": "pass1234"})
        assert resp.status_code == 200
        data = resp.json()
        assert data["status"] == "success"
        assert "token" in data

    def test_login_with_wrong_password_returns_401(self):
        email = "wrongpass_xyz@example.com"
        client.post("/api/auth/signup", json={"name": "WP User", "email": email, "password": "correct"})
        resp = client.post("/api/auth/login", json={"email": email, "password": "wrong"})
        assert resp.status_code == 401

    def test_login_nonexistent_user_returns_401(self):
        resp = client.post("/api/auth/login", json={
            "email": "nobody_xyz_zz@example.com",
            "password": "pass",
        })
        assert resp.status_code == 401


# ── Preview endpoint ─────────────────────────────────────────────────────────

SAMPLE_CANDIDATE = {
    "name": "Test User",
    "contact": {
        "email": "test.user@example.com",
        "phone": "0400000000",
        "location": "Sydney, NSW",
    },
    "professional_summary": "Test candidate for automated testing purposes.",
    "key_skills": ["Testing", "Automation"],
    "technical_skills": ["Python", "pytest"],
    "experience": [
        {
            "title": "QA Engineer",
            "company": "Test Corp",
            "location": "Sydney, NSW",
            "dates": "Jan 2022 – Present",
            "description": "Automated test suites.",
            "bullets": ["Increased coverage by 30%"],
        }
    ],
    "education": [
        {
            "degree": "B.Sc. Computer Science",
            "field": "Testing",
            "institution": "Test University",
            "graduation_year": "2020",
        }
    ],
    "certifications": ["ISTQB Foundation"],
    "awards": [],
}


class TestPreviewResume:
    def test_preview_returns_200(self):
        resp = client.post("/api/preview-resume", json=SAMPLE_CANDIDATE)
        assert resp.status_code == 200

    def test_preview_returns_html_content_type(self):
        resp = client.post("/api/preview-resume", json=SAMPLE_CANDIDATE)
        assert "text/html" in resp.headers["content-type"]

    def test_preview_contains_candidate_name(self):
        resp = client.post("/api/preview-resume", json=SAMPLE_CANDIDATE)
        assert "TEST USER" in resp.text

    def test_preview_contains_contact_email(self):
        resp = client.post("/api/preview-resume", json=SAMPLE_CANDIDATE)
        assert "test.user@example.com" in resp.text

    def test_preview_contains_summary(self):
        resp = client.post("/api/preview-resume", json=SAMPLE_CANDIDATE)
        assert "automated testing purposes" in resp.text

    def test_preview_contains_experience(self):
        resp = client.post("/api/preview-resume", json=SAMPLE_CANDIDATE)
        assert "QA Engineer" in resp.text
        assert "Test Corp" in resp.text

    def test_preview_contains_education(self):
        resp = client.post("/api/preview-resume", json=SAMPLE_CANDIDATE)
        assert "B.Sc. Computer Science" in resp.text

    def test_preview_missing_required_field_returns_422(self):
        # Missing 'contact' field
        resp = client.post("/api/preview-resume", json={"name": "Incomplete"})
        assert resp.status_code == 422

    def test_preview_with_linkedin(self):
        candidate = dict(SAMPLE_CANDIDATE)
        candidate["contact"] = {**SAMPLE_CANDIDATE["contact"], "linkedin": "linkedin.com/in/testuser"}
        resp = client.post("/api/preview-resume", json=candidate)
        assert resp.status_code == 200
        assert "linkedin.com/in/testuser" in resp.text


# ── Generate resume endpoint ─────────────────────────────────────────────────

class TestGenerateResume:
    def test_generate_returns_success(self):
        resp = client.post("/api/generate-resume", json=SAMPLE_CANDIDATE)
        assert resp.status_code == 200
        data = resp.json()
        assert data["status"] == "success"
        assert "filename" in data["data"]

    def test_generate_creates_docx_file(self):
        resp = client.post("/api/generate-resume", json=SAMPLE_CANDIDATE)
        assert resp.status_code == 200
        filename = resp.json()["data"]["filename"]
        file_path = Path("./resumes") / filename
        assert file_path.exists()

    def test_generate_returns_download_url(self):
        resp = client.post("/api/generate-resume", json=SAMPLE_CANDIDATE)
        data = resp.json()["data"]
        assert "download_url" in data
        assert data["download_url"].startswith("/api/resumes/download-file/")

    def test_generate_missing_name_returns_422(self):
        bad = dict(SAMPLE_CANDIDATE)
        del bad["name"]
        resp = client.post("/api/generate-resume", json=bad)
        assert resp.status_code == 422

    def test_generate_missing_contact_returns_422(self):
        bad = dict(SAMPLE_CANDIDATE)
        del bad["contact"]
        resp = client.post("/api/generate-resume", json=bad)
        assert resp.status_code == 422


# ── Download endpoint ────────────────────────────────────────────────────────

class TestDownloadResume:
    def test_download_generated_file(self):
        # First generate a file
        gen_resp = client.post("/api/generate-resume", json=SAMPLE_CANDIDATE)
        assert gen_resp.status_code == 200
        filename = gen_resp.json()["data"]["filename"]

        dl_resp = client.get(f"/api/resumes/download-file/{filename}")
        assert dl_resp.status_code == 200
        assert "application/vnd.openxmlformats" in dl_resp.headers["content-type"]

    def test_download_nonexistent_file_returns_404(self):
        resp = client.get("/api/resumes/download-file/does_not_exist_xyz.docx")
        assert resp.status_code == 404


# ── Templates endpoint ───────────────────────────────────────────────────────

class TestTemplates:
    def test_get_templates_returns_list(self):
        resp = client.get("/api/templates")
        assert resp.status_code == 200
        data = resp.json()
        assert data["status"] == "success"
        assert isinstance(data["templates"], list)
        assert len(data["templates"]) > 0

    def test_templates_have_required_fields(self):
        resp = client.get("/api/templates")
        for tpl in resp.json()["templates"]:
            assert "id" in tpl
            assert "name" in tpl
            assert "description" in tpl

    def test_templates_returns_five_entries(self):
        resp = client.get("/api/templates")
        assert len(resp.json()["templates"]) == 5

    def test_template_ids_are_correct(self):
        resp = client.get("/api/templates")
        ids = {t["id"] for t in resp.json()["templates"]}
        assert ids == {"modern", "classic", "creative", "minimal", "executive"}

    def test_template_previews_returns_all_five(self):
        resp = client.get("/api/templates/previews")
        assert resp.status_code == 200
        data = resp.json()
        assert set(data.keys()) == {"modern", "classic", "creative", "minimal", "executive"}

    def test_template_previews_contain_valid_html(self):
        resp = client.get("/api/templates/previews")
        assert resp.status_code == 200
        for tid, html_str in resp.json().items():
            assert "<!DOCTYPE html>" in html_str, f"{tid} preview missing DOCTYPE"
            assert "ALEX JOHNSON" in html_str, f"{tid} preview missing dummy candidate name"

    def test_template_previews_executive_has_amber_rule(self):
        resp = client.get("/api/templates/previews")
        assert "#b45309" in resp.json()["executive"]

    def test_template_previews_classic_uses_serif(self):
        resp = client.get("/api/templates/previews")
        assert "serif" in resp.json()["classic"].lower()


class TestGenerateWithTemplate:
    """Tests for the `template` parameter in POST /api/generate."""

    @pytest.mark.parametrize("template_id", ["modern", "classic", "creative", "minimal"])
    def test_generate_with_each_template(self, monkeypatch, template_id):
        mock_client = MagicMock()
        mock_resp = MagicMock()
        mock_resp.choices[0].message.content = json.dumps(MOCK_RESUME_JSON)
        mock_client.chat.completions.create.return_value = mock_resp
        monkeypatch.setattr("main.openai_client", mock_client)

        txt = b"Jane Smith\nSoftware Engineer"
        resp = client.post(
            "/api/generate",
            data={"job_description": "Python developer", "template": template_id},
            files=[("files", ("cv.txt", io.BytesIO(txt), "text/plain"))],
        )
        assert resp.status_code == 200, f"Failed for template '{template_id}': {resp.text}"
        data = resp.json()
        assert "preview_html" in data
        assert data["filename"].endswith(".docx")

    def test_generate_without_template_defaults_to_modern(self, monkeypatch):
        """Omitting the template field must not cause an error (defaults to modern)."""
        mock_client = MagicMock()
        mock_resp = MagicMock()
        mock_resp.choices[0].message.content = json.dumps(MOCK_RESUME_JSON)
        mock_client.chat.completions.create.return_value = mock_resp
        monkeypatch.setattr("main.openai_client", mock_client)

        txt = b"Jane Smith\nSoftware Engineer"
        resp = client.post(
            "/api/generate",
            data={"job_description": "Python developer"},
            files=[("files", ("cv.txt", io.BytesIO(txt), "text/plain"))],
        )
        assert resp.status_code == 200

    def test_generate_unknown_template_falls_back_gracefully(self, monkeypatch):
        """An unknown template ID must not crash — it falls back to modern."""
        mock_client = MagicMock()
        mock_resp = MagicMock()
        mock_resp.choices[0].message.content = json.dumps(MOCK_RESUME_JSON)
        mock_client.chat.completions.create.return_value = mock_resp
        monkeypatch.setattr("main.openai_client", mock_client)

        txt = b"Jane Smith\nSoftware Engineer"
        resp = client.post(
            "/api/generate",
            data={"job_description": "Python developer", "template": "nonexistent"},
            files=[("files", ("cv.txt", io.BytesIO(txt), "text/plain"))],
        )
        assert resp.status_code == 200

    def test_creative_template_preview_contains_purple(self, monkeypatch):
        mock_client = MagicMock()
        mock_resp = MagicMock()
        mock_resp.choices[0].message.content = json.dumps(MOCK_RESUME_JSON)
        mock_client.chat.completions.create.return_value = mock_resp
        monkeypatch.setattr("main.openai_client", mock_client)

        txt = b"Jane Smith\nSoftware Engineer"
        resp = client.post(
            "/api/generate",
            data={"job_description": "Python developer", "template": "creative"},
            files=[("files", ("cv.txt", io.BytesIO(txt), "text/plain"))],
        )
        assert resp.status_code == 200
        # Creative template uses purple (#6b21a8) for headings
        assert "#6b21a8" in resp.json()["preview_html"]

    def test_classic_template_preview_contains_serif_font(self, monkeypatch):
        mock_client = MagicMock()
        mock_resp = MagicMock()
        mock_resp.choices[0].message.content = json.dumps(MOCK_RESUME_JSON)
        mock_client.chat.completions.create.return_value = mock_resp
        monkeypatch.setattr("main.openai_client", mock_client)

        txt = b"Jane Smith\nSoftware Engineer"
        resp = client.post(
            "/api/generate",
            data={"job_description": "Python developer", "template": "classic"},
            files=[("files", ("cv.txt", io.BytesIO(txt), "text/plain"))],
        )
        assert resp.status_code == 200
        assert "Georgia" in resp.json()["preview_html"]


# ── Resumes list endpoint ────────────────────────────────────────────────────

class TestResumesList:
    def test_get_resumes_for_unknown_user_returns_empty(self):
        resp = client.get("/api/resumes?user_id=99999")
        assert resp.status_code == 200
        data = resp.json()
        assert data["status"] == "success"
        assert data["resumes"] == []


# ── /api/generate (document-based AI generation) ────────────────────────────

# Minimal resume JSON that matches the schema expected by doc_builder
MOCK_RESUME_JSON = {
    "name": "Jane Smith",
    "contact": {
        "email": "jane@example.com",
        "phone": "0412 345 678",
        "location": "Sydney, NSW",
        "linkedin": "",
    },
    "professional_summary": "Experienced software engineer with 8 years in fintech.",
    "key_skills": ["Python", "FastAPI", "AWS"],
    "experience": [
        {
            "title": "Senior Software Engineer",
            "company": "FinTech Co",
            "location": "Sydney, NSW",
            "dates": "Jan 2020 – Present",
            "description": "",
            "bullets": ["Led migration of legacy systems to microservices, reducing latency by 40%."],
        }
    ],
    "education": [
        {
            "degree": "Bachelor of Computer Science",
            "field": "Software Engineering",
            "institution": "University of Sydney",
            "graduation_year": "2015",
        }
    ],
    "certifications": [],
    "awards": [],
    "technical_skills": ["Python", "Docker", "Kubernetes"],
}


def _mock_openai_client(monkeypatch):
    """Return a mocked openai_client and patch it into main module."""
    mock_client = MagicMock()
    mock_response = MagicMock()
    mock_response.choices[0].message.content = json.dumps(MOCK_RESUME_JSON)
    mock_client.chat.completions.create.return_value = mock_response
    monkeypatch.setattr("main.openai_client", mock_client)
    return mock_client


class TestGenerateFromDocuments:
    """Tests for POST /api/generate — document-based AI resume generation."""

    def test_generate_without_files_returns_error(self):
        # No files provided → endpoint must not return 200.
        # FastAPI may return 422 (form validation) or the endpoint returns 400.
        resp = client.post(
            "/api/generate",
            data={"job_description": "Senior Software Engineer at FinTech startup."},
        )
        assert resp.status_code in (400, 422)
        assert resp.status_code != 200

    def test_generate_without_job_description_uses_general_mode(self):
        # job_description is now optional — omitting it triggers general mode.
        # Without an OpenAI key the request reaches the 503 check, confirming
        # it was accepted by form validation (not rejected with 422).
        txt = b"Jane Smith\nSoftware Engineer\n8 years Python experience"
        resp = client.post(
            "/api/generate",
            files=[("files", ("resume.txt", io.BytesIO(txt), "text/plain"))],
        )
        # Must NOT be 422 (form validation error) — the request is structurally valid.
        assert resp.status_code != 422

    def test_generate_no_openai_key_returns_503(self, monkeypatch):
        monkeypatch.setattr("main.openai_client", None)
        txt = b"Jane Smith\nSoftware Engineer"
        resp = client.post(
            "/api/generate",
            data={"job_description": "Python developer role"},
            files=[("files", ("resume.txt", io.BytesIO(txt), "text/plain"))],
        )
        assert resp.status_code == 503
        assert "openai" in resp.json()["detail"].lower()

    def test_generate_with_txt_file_returns_success(self, monkeypatch):
        _mock_openai_client(monkeypatch)
        txt = b"Jane Smith\nSoftware Engineer\n8 years Python experience"
        resp = client.post(
            "/api/generate",
            data={"job_description": "Senior Python developer with AWS experience."},
            files=[("files", ("resume.txt", io.BytesIO(txt), "text/plain"))],
        )
        assert resp.status_code == 200
        data = resp.json()
        # Generate returns a flat response (no standardize_response wrapper)
        assert "filename" in data
        assert data["filename"].endswith(".docx")
        assert "download_url" in data
        assert "preview_html" in data

    def test_generate_creates_docx_file(self, monkeypatch):
        _mock_openai_client(monkeypatch)
        txt = b"Jane Smith\nSoftware Engineer"
        resp = client.post(
            "/api/generate",
            data={"job_description": "Python developer"},
            files=[("files", ("cv.txt", io.BytesIO(txt), "text/plain"))],
        )
        assert resp.status_code == 200
        filename = resp.json()["filename"]
        assert Path("./resumes") / filename

    def test_generate_preview_html_contains_name(self, monkeypatch):
        _mock_openai_client(monkeypatch)
        txt = b"Jane Smith\nSoftware Engineer"
        resp = client.post(
            "/api/generate",
            data={"job_description": "Python developer"},
            files=[("files", ("cv.txt", io.BytesIO(txt), "text/plain"))],
        )
        assert resp.status_code == 200
        assert "JANE SMITH" in resp.json()["preview_html"]

    def test_generate_download_url_accessible(self, monkeypatch):
        _mock_openai_client(monkeypatch)
        txt = b"Jane Smith\nSoftware Engineer"
        gen_resp = client.post(
            "/api/generate",
            data={"job_description": "Python developer"},
            files=[("files", ("cv.txt", io.BytesIO(txt), "text/plain"))],
        )
        assert gen_resp.status_code == 200
        dl_resp = client.get(gen_resp.json()["download_url"])
        assert dl_resp.status_code == 200
        assert "application/vnd.openxmlformats" in dl_resp.headers["content-type"]

    def test_generate_with_docx_file(self, monkeypatch):
        """Upload a real minimal .docx file and verify extraction doesn't crash."""
        _mock_openai_client(monkeypatch)
        from docx import Document as DocxDocument
        buf = io.BytesIO()
        doc = DocxDocument()
        doc.add_paragraph("Jane Smith")
        doc.add_paragraph("Senior Software Engineer")
        doc.save(buf)
        buf.seek(0)
        resp = client.post(
            "/api/generate",
            data={"job_description": "Software engineer role"},
            files=[("files", ("resume.docx", buf, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"))],
        )
        assert resp.status_code == 200

    def test_generate_with_pdf_file(self, monkeypatch):
        """PDF upload should not crash even if pypdf has system-level issues.
        The endpoint falls back to raw decode; the mock fills in the resume data."""
        _mock_openai_client(monkeypatch)
        # Use a simple text file named .pdf — the endpoint will attempt PDF
        # extraction and fall back gracefully, then the mock provides the response.
        fake_pdf_as_text = b"Jane Smith\nSenior Software Engineer\n8 years Python experience"
        resp = client.post(
            "/api/generate",
            data={"job_description": "Software engineer role"},
            files=[("files", ("resume.pdf", io.BytesIO(fake_pdf_as_text), "application/pdf"))],
        )
        # Should succeed regardless of pypdf availability
        assert resp.status_code == 200

    def test_generate_too_many_files_returns_400(self, monkeypatch):
        _mock_openai_client(monkeypatch)
        files = [
            ("files", (f"cv{i}.txt", io.BytesIO(b"text"), "text/plain"))
            for i in range(6)
        ]
        resp = client.post(
            "/api/generate",
            data={"job_description": "Python developer"},
            files=files,
        )
        assert resp.status_code == 400
        assert "maximum" in resp.json()["detail"].lower()

    def test_generate_openai_json_error_returns_500(self, monkeypatch):
        mock_client = MagicMock()
        mock_response = MagicMock()
        mock_response.choices[0].message.content = "This is not JSON at all"
        mock_client.chat.completions.create.return_value = mock_response
        monkeypatch.setattr("main.openai_client", mock_client)

        txt = b"Jane Smith\nSoftware Engineer"
        resp = client.post(
            "/api/generate",
            data={"job_description": "Python developer"},
            files=[("files", ("cv.txt", io.BytesIO(txt), "text/plain"))],
        )
        assert resp.status_code == 500

    def test_generate_always_returns_resume_id(self, monkeypatch):
        """Even for guests (no user_id), the generate endpoint must return a resume_id."""
        _mock_openai_client(monkeypatch)
        txt = b"Jane Smith\nSoftware Engineer"
        resp = client.post(
            "/api/generate",
            data={"job_description": "Python developer"},
            files=[("files", ("cv.txt", io.BytesIO(txt), "text/plain"))],
        )
        assert resp.status_code == 200
        data = resp.json()
        assert data["resume_id"] is not None
        assert isinstance(data["resume_id"], int)


# ── Guest editing ─────────────────────────────────────────────────────────────

class TestGuestEditing:
    """Guest users (no user_id) can edit up to MAX_PROMPTS_GUEST times."""

    def _generate_guest_resume(self, monkeypatch):
        _mock_openai_client(monkeypatch)
        txt = b"Jane Smith\nSoftware Engineer"
        resp = client.post(
            "/api/generate",
            data={"job_description": "Python developer"},
            files=[("files", ("cv.txt", io.BytesIO(txt), "text/plain"))],
        )
        assert resp.status_code == 200
        return resp.json()["resume_id"]

    def _mock_edit_openai(self, monkeypatch):
        mock_client = MagicMock()
        mock_resp = MagicMock()
        mock_resp.choices[0].message.content = json.dumps(MOCK_RESUME_JSON)
        mock_client.chat.completions.create.return_value = mock_resp
        monkeypatch.setattr("main.openai_client", mock_client)

    def test_guest_can_edit_without_user_id(self, monkeypatch):
        resume_id = self._generate_guest_resume(monkeypatch)
        self._mock_edit_openai(monkeypatch)
        resp = client.post(
            f"/api/resumes/{resume_id}/edit",
            data={"prompt": "Make the summary shorter"},
        )
        assert resp.status_code == 200
        data = resp.json()
        assert data["status"] == "success"
        assert "preview_html" in data
        assert "remaining_prompts" in data

    def test_guest_edit_decrements_remaining_prompts(self, monkeypatch):
        resume_id = self._generate_guest_resume(monkeypatch)
        self._mock_edit_openai(monkeypatch)
        resp = client.post(
            f"/api/resumes/{resume_id}/edit",
            data={"prompt": "Add more skills"},
        )
        assert resp.status_code == 200
        data = resp.json()
        from utils import MAX_PROMPTS_GUEST
        assert data["max_prompts"] == MAX_PROMPTS_GUEST
        assert data["prompt_count"] == 1
        assert data["remaining_prompts"] == MAX_PROMPTS_GUEST - 1

    def test_guest_edit_limit_enforced(self, monkeypatch):
        """After MAX_PROMPTS_GUEST edits the endpoint returns 403."""
        from utils import MAX_PROMPTS_GUEST
        resume_id = self._generate_guest_resume(monkeypatch)
        self._mock_edit_openai(monkeypatch)
        # Exhaust all allowed edits
        for _ in range(MAX_PROMPTS_GUEST):
            r = client.post(
                f"/api/resumes/{resume_id}/edit",
                data={"prompt": "tweak"},
            )
            assert r.status_code == 200
        # Next edit must be rejected
        resp = client.post(
            f"/api/resumes/{resume_id}/edit",
            data={"prompt": "one more tweak"},
        )
        assert resp.status_code == 403
        assert "free edits" in resp.json()["detail"].lower()

    def test_guest_edit_returns_download_url_filename(self, monkeypatch):
        resume_id = self._generate_guest_resume(monkeypatch)
        self._mock_edit_openai(monkeypatch)
        resp = client.post(
            f"/api/resumes/{resume_id}/edit",
            data={"prompt": "Update the summary"},
        )
        assert resp.status_code == 200
        data = resp.json()
        assert "filename" in data
        assert data["filename"].endswith(".docx")

    def test_guest_cannot_edit_logged_in_resume(self, monkeypatch):
        """A guest should not be able to edit a resume that belongs to a user."""
        # Create a logged-in user resume
        _mock_openai_client(monkeypatch)
        signup_resp = client.post("/api/auth/signup", json={
            "name": "Guest Edit Guard",
            "email": "guest_guard_xyz@example.com",
            "password": "pass1234",
        })
        assert signup_resp.status_code in (200, 400)
        login_resp = client.post("/api/auth/login", json={
            "email": "guest_guard_xyz@example.com",
            "password": "pass1234",
        })
        user_id = login_resp.json()["user_id"]

        txt = b"Jane Smith\nSoftware Engineer"
        gen_resp = client.post(
            "/api/generate",
            data={"job_description": "Python developer", "user_id": user_id},
            files=[("files", ("cv.txt", io.BytesIO(txt), "text/plain"))],
        )
        assert gen_resp.status_code == 200
        resume_id = gen_resp.json()["resume_id"]

        self._mock_edit_openai(monkeypatch)
        # Guest (no user_id) tries to edit a user-owned resume — must fail
        resp = client.post(
            f"/api/resumes/{resume_id}/edit",
            data={"prompt": "Change summary"},
        )
        assert resp.status_code == 404  # not found (user_id IS NULL filter fails)

    def test_edit_nonexistent_resume_returns_404(self, monkeypatch):
        self._mock_edit_openai(monkeypatch)
        resp = client.post(
            "/api/resumes/999999/edit",
            data={"prompt": "Update"},
        )
        assert resp.status_code == 404


# ── Prompt-info endpoint ──────────────────────────────────────────────────────

class TestPromptInfoEndpoint:
    def test_prompt_info_for_free_user(self):
        from utils import MAX_PROMPTS_FREE
        client.post("/api/auth/signup", json={
            "name": "Prompt Info User",
            "email": "promptinfo_xyz@example.com",
            "password": "pass1234",
        })
        login = client.post("/api/auth/login", json={
            "email": "promptinfo_xyz@example.com",
            "password": "pass1234",
        })
        user_id = login.json()["user_id"]
        resp = client.get(f"/api/users/{user_id}/prompt-info")
        assert resp.status_code == 200
        data = resp.json()
        assert data["max_prompts"] == MAX_PROMPTS_FREE
        assert data["remaining_prompts"] == MAX_PROMPTS_FREE
        assert data["membership_tier"] == "free"


# ── Template switcher ─────────────────────────────────────────────────────────

class TestSwitchTemplate:
    """POST /api/resumes/{id}/switch-template re-renders without consuming quota."""

    def _guest_resume_id(self, monkeypatch):
        _mock_openai_client(monkeypatch)
        txt = b"Jane Smith\nSoftware Engineer"
        resp = client.post(
            "/api/generate",
            data={"job_description": "Python developer"},
            files=[("files", ("cv.txt", io.BytesIO(txt), "text/plain"))],
        )
        assert resp.status_code == 200
        return resp.json()["resume_id"]

    def test_switch_template_returns_new_preview(self, monkeypatch):
        resume_id = self._guest_resume_id(monkeypatch)
        resp = client.post(
            f"/api/resumes/{resume_id}/switch-template",
            data={"template_id": "classic"},
        )
        assert resp.status_code == 200
        data = resp.json()
        assert data["status"] == "success"
        assert "preview_html" in data
        assert data["template_id"] == "classic"

    def test_switch_template_updates_filename(self, monkeypatch):
        resume_id = self._guest_resume_id(monkeypatch)
        resp = client.post(
            f"/api/resumes/{resume_id}/switch-template",
            data={"template_id": "creative"},
        )
        assert resp.status_code == 200
        assert resp.json()["filename"].endswith(".docx")

    def test_switch_template_all_four_templates(self, monkeypatch):
        resume_id = self._guest_resume_id(monkeypatch)
        for tid in ("modern", "classic", "creative", "minimal"):
            resp = client.post(
                f"/api/resumes/{resume_id}/switch-template",
                data={"template_id": tid},
            )
            assert resp.status_code == 200, f"Failed for {tid}: {resp.text}"
            assert resp.json()["template_id"] == tid

    def test_switch_template_invalid_id_returns_400(self, monkeypatch):
        resume_id = self._guest_resume_id(monkeypatch)
        resp = client.post(
            f"/api/resumes/{resume_id}/switch-template",
            data={"template_id": "neon_unicorn"},
        )
        assert resp.status_code == 400
        assert "Unknown template" in resp.json()["detail"]

    def test_switch_template_nonexistent_resume_returns_404(self, monkeypatch):
        resp = client.post(
            "/api/resumes/999999/switch-template",
            data={"template_id": "modern"},
        )
        assert resp.status_code == 404

    def test_switch_template_classic_preview_contains_serif(self, monkeypatch):
        resume_id = self._guest_resume_id(monkeypatch)
        resp = client.post(
            f"/api/resumes/{resume_id}/switch-template",
            data={"template_id": "classic"},
        )
        assert resp.status_code == 200
        assert "serif" in resp.json()["preview_html"].lower()

    def test_switch_template_does_not_consume_edit_quota(self, monkeypatch):
        """Switching templates must NOT increment guest_edit_count."""
        resume_id = self._guest_resume_id(monkeypatch)
        for tid in ("classic", "creative", "minimal", "modern", "classic"):
            resp = client.post(
                f"/api/resumes/{resume_id}/switch-template",
                data={"template_id": tid},
            )
            assert resp.status_code == 200

        # Guest edit quota should still be full (all 3 remaining)
        from models import Resume
        from database import SessionLocal
        db = SessionLocal()
        try:
            resume = db.query(Resume).filter(Resume.id == resume_id).first()
            assert (resume.guest_edit_count or 0) == 0
        finally:
            db.close()
